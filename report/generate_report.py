# -*- coding: utf-8 -*-
"""用 python-docx 生成毕业论文格式的实训报告（两个项目合一）。

结构：封面 → 摘要/关键词 → 第1章 引言 → 第2章 相关理论基础 →
第3章 项目一(红酒) → 第4章 项目二(AlexNet) → 第5章 结论与实验总结 → 参考文献。
图表数据来自两个任务 outputs/ 目录。
"""
import os
import csv
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------- OMML 公式构建（Word 公式编辑器格式） ----------------------
_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _esc(t):
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def mr(t):
    """公式中的普通文本片段。"""
    return f'<m:r><m:t xml:space="preserve">{_esc(t)}</m:t></m:r>'


def mfrac(num, den):
    """分式。"""
    return f"<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>"


def msup(base, sup):
    """上标。"""
    return f"<m:sSup><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>"


def msub(base, sub):
    """下标。"""
    return f"<m:sSub><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>"


def msubsup(base, sub, sup):
    """同时带上下标。"""
    return (f"<m:sSubSup><m:e>{base}</m:e><m:sub>{sub}</m:sub>"
            f"<m:sup>{sup}</m:sup></m:sSubSup>")

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
HERE = os.path.dirname(os.path.abspath(__file__))
T1 = os.path.join(ROOT, "task1_ml_wine", "outputs")
T2 = os.path.join(ROOT, "task2_alexnet_fmnist", "outputs")


# ---------------------- 辅助函数 ----------------------
def set_base_style(doc):
    """设置正文中文字体（宋体）与字号（小四=12pt）。"""
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "SimSun")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def para(doc, text, indent=True, align=None):
    """添加正文段落，默认首行缩进 2 字符。"""
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    if align is not None:
        p.alignment = align
    return p


def equation(doc, inner):
    """以居中段落插入一个 OMML 公式（Word 公式编辑器格式）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.append(parse_xml(f'<m:oMath xmlns:m="{_M}">{inner}</m:oMath>'))
    return p


def add_image(doc, path, width=5.4, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].font.size = Pt(10.5)


def add_table(doc, header, rows, caption=None):
    if caption:
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].font.size = Pt(10.5)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light List Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(header):
        table.rows[0].cells[j].text = str(h)
    for r in rows:
        cells = table.add_row().cells
        for j, v in enumerate(r):
            cells[j].text = str(v)
    return table


# ---------------------- 封面与摘要 ----------------------
def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph("人工智能综合实训 II"); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(28); t.runs[0].font.bold = True
    s = doc.add_paragraph("实 训 报 告"); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(22); s.runs[0].font.bold = True
    for _ in range(3):
        doc.add_paragraph()
    info = [("题　　目：", "机器学习与手写 AlexNet 图像分类"),
            ("学　　院：", "人工智能与低空技术学院"),
            ("专　　业：", "人工智能"),
            ("小组成员：", "雷正　蔡铭飞　冼嘉谦"),
            ("指导教师：", "赵静")]
    for k, v in info:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(k + v); run.font.size = Pt(14)
    for _ in range(4):
        doc.add_paragraph()
    d = doc.add_paragraph("2026 年 5 月"); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.runs[0].font.size = Pt(14)
    doc.add_page_break()


def abstract(doc):
    h = doc.add_paragraph("摘　要"); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.size = Pt(16); h.runs[0].font.bold = True
    para(doc,
         "随着人工智能技术的迅速发展，机器学习已成为数据分析与智能决策的核心工具。本实训围绕"
         "“机器学习全流程实践”展开，完成两个相互独立又层层递进的项目，系统覆盖结构化数据建模与"
         "图像深度学习两大方向。项目一面向结构化数据，在 UCI 白葡萄酒质量数据集上构建多分类模型，"
         "系统对比了支持向量机、决策树、随机森林与逻辑回归四种算法，采用网格搜索与五折交叉验证进行"
         "超参数寻优，并使用准确率、精确率、召回率、宏平均 F1 等指标评估。实验表明随机森林取得最佳"
         "性能，测试集准确率 73.5%、宏平均 F1 为 0.734，显著优于线性模型。项目二面向图像数据，依据"
         "经典论文逐层手工实现卷积神经网络 AlexNet（不调用任何预置模型），在 Fashion-MNIST 数据集上"
         "完成十分类任务，经 15 轮训练后测试集准确率达到 91.55%、宏平均 F1 为 0.915。两个项目完整覆盖"
         "了数据加载、可视化、预处理、模型建立、模型评估与结论分析的全流程，验证了集成学习在结构化"
         "数据上的优势以及深度卷积网络在图像识别上的有效性，并提出了若干改进方向。")
    para(doc, "关键词：机器学习；支持向量机；随机森林；卷积神经网络；AlexNet；图像分类")
    doc.add_page_break()


# ---------------------- 第1章 引言 ----------------------
def chapter1(doc):
    doc.add_heading("第1章 引言", level=1)
    doc.add_heading("1.1 研究背景与意义", level=2)
    para(doc, "人工智能是研究、开发用于模拟和扩展人类智能的理论、方法与技术的科学，而机器学习作为其"
              "最重要的分支，致力于让计算机系统通过经验（数据）自动改善性能，而无需被显式编程。近年来，"
              "得益于数据规模的爆炸式增长与计算硬件（尤其是 GPU）的进步，机器学习在图像识别、自然语言"
              "处理、推荐系统、医疗诊断等众多领域取得了突破性进展，深刻改变了人们的生产与生活方式。")
    para(doc, "按照数据形态的不同，机器学习任务大致可分为两类：一类是面向结构化（表格）数据的任务，"
              "典型如分类与回归，通常采用支持向量机、决策树、集成学习等传统算法；另一类是面向图像、语音、"
              "文本等非结构化数据的任务，主要依赖深度神经网络。本实训分别选取这两类任务的代表性问题——"
              "红酒质量分级与服饰图像分类，通过完整的工程实践，加深对机器学习核心思想、典型算法与建模"
              "流程的理解，培养从数据到模型、再到结论的综合实践能力，具有重要的学习与应用价值。")
    doc.add_heading("1.2 国内外研究现状", level=2)
    para(doc, "在传统机器学习方面，支持向量机由 Vapnik 等人于 20 世纪 90 年代提出，凭借坚实的统计学习"
              "理论基础和优秀的小样本泛化能力，长期是分类任务的重要工具；以随机森林、梯度提升树（GBDT、"
              "XGBoost、LightGBM）为代表的集成学习方法，通过组合多个弱学习器进一步提升了精度与鲁棒性，"
              "至今仍是结构化数据竞赛中的主流方案。在深度学习方面，2012 年 Krizhevsky 等人提出的 AlexNet"
              "在 ImageNet 竞赛中以显著优势夺冠，掀起了深度卷积网络的研究热潮；此后 VGG、GoogLeNet、"
              "ResNet 等更深、更高效的网络结构相继涌现，不断刷新图像识别的精度上限。本实训以经典且最具"
              "代表性的 AlexNet 为研究对象，逐层复现其结构，旨在透彻理解卷积神经网络的工作机制。")
    doc.add_heading("1.3 项目研究思路", level=2)
    para(doc, "本实训采用统一的机器学习方法论，对两个项目均遵循以下六个环节循序推进，整体流程见图 1-1：")
    for line in [
        "（1）数据加载：从公开数据源获取数据集，读入内存并核对规模、字段与完整性；",
        "（2）数据查看：通过描述性统计与可视化了解数据分布、相关性与潜在问题；",
        "（3）数据预处理：包括标签构造、特征标准化、数据集划分以及类别不平衡处理；",
        "（4）模型建立：根据任务特点选择合适的算法或网络结构，并通过调参确定最优配置；",
        "（5）模型评估：在独立测试集上用多种指标量化性能，并借助混淆矩阵等工具深入分析；",
        "（6）结论分析：归纳实验结果，对比模型优劣，总结经验并提出改进方案。"]:
        para(doc, line, indent=False)
    add_image(doc, os.path.join(HERE, "flowchart.png"), 4.8, "图 1-1 机器学习项目统一研究流程")
    doc.add_heading("1.4 小组成员分工", level=2)
    para(doc, "本项目由小组协作完成，成员分工如表 1-1 所示。")
    add_table(doc, ["学号", "姓名", "主要分工"],
              [["202434610309", "雷正", "项目一数据预处理与机器学习建模、调参与评估"],
               ["202434610301", "蔡铭飞", "项目二 AlexNet 网络实现、模型训练与可视化评估"],
               ["202434610326", "冼嘉谦", "数据可视化、报告撰写、PPT 制作与汇报"]],
              caption="表 1-1 小组成员分工表")
    doc.add_page_break()


# ---------------------- 第2章 相关理论基础 ----------------------
def chapter2(doc):
    doc.add_heading("第2章 相关理论基础", level=1)
    doc.add_heading("2.1 监督学习与分类问题", level=2)
    para(doc, "监督学习是机器学习中最常见的范式，其任务是在给定带标签的训练集 D={(xᵢ, yᵢ)} 上学习一个"
              "映射 f: X→Y，使其对未知样本具有良好的预测能力。当标签 y 取离散值时即为分类问题。本实训"
              "两个项目均属于多分类问题：项目一为三分类（红酒质量等级），项目二为十分类（服饰类别）。"
              "模型学习的核心是最小化训练数据上的损失函数，同时通过正则化、交叉验证等手段控制模型复杂度"
              "以提升泛化能力。")

    doc.add_heading("2.2 传统机器学习算法", level=2)
    doc.add_heading("2.2.1 支持向量机（SVM）", level=3)
    para(doc, "支持向量机的基本思想是在特征空间中寻找一个最大间隔超平面以分隔不同类别，在约束 "
              "yᵢ(wᵀxᵢ+b) ≥ 1−ξᵢ 下，其优化目标为：")
    equation(doc, mr("min  ") + mfrac(mr("1"), mr("2")) + msup(mr("‖w‖"), mr("2"))
             + mr(" + C ") + mr("Σ") + msub(mr("ξ"), mr("i")))
    para(doc, "其中 C 为惩罚系数、ξᵢ 为松弛变量。对于非线性可分问题，SVM 通过核函数将样本映射到高维"
              "空间，本实训采用径向基（RBF）核：")
    equation(doc, mr("K(x, x′) = exp(−γ") + msup(mr("‖x−x′‖"), mr("2")) + mr(")"))

    doc.add_heading("2.2.2 决策树", level=3)
    para(doc, "决策树通过递归地选择最优特征对样本空间进行划分，形成树状的判别结构。划分依据通常为基尼"
              "不纯度：")
    equation(doc, mr("Gini(D) = 1 − ") + mr("Σ") + msubsup(mr("p"), mr("k"), mr("2")))
    para(doc, "其中 pₖ 为类别 k 在节点中的占比。算法每次选择使子节点不纯度下降最大的特征进行分裂，"
              "具有良好的可解释性，但单棵树容易过拟合。")

    doc.add_heading("2.2.3 随机森林", level=3)
    para(doc, "随机森林是一种基于 Bagging 思想的集成学习方法，通过自助采样（Bootstrap）构建多棵决策树，"
              "并在每次分裂时随机选取部分特征，最后对所有树的预测结果进行多数表决。随机性的引入降低了"
              "模型方差，使其在保持高精度的同时具有较强的抗过拟合能力。")

    doc.add_heading("2.2.4 逻辑回归", level=3)
    para(doc, "逻辑回归是一种线性分类模型，通过 Softmax（多分类）函数将线性输出映射为类别概率：")
    equation(doc, mr("P(y=k|x) = ")
             + mfrac(msup(mr("e"), msub(mr("w"), mr("k")) + mr("ᵀx")),
                     mr("Σ") + msup(mr("e"), msub(mr("w"), mr("j")) + mr("ᵀx"))))
    para(doc, "并以交叉熵为损失函数进行参数估计。其结构简单、训练高效，常作为分类任务的基线模型。")

    doc.add_heading("2.3 卷积神经网络与 AlexNet", level=2)
    para(doc, "卷积神经网络（CNN）是处理图像数据的主流深度模型，其核心组件包括：")
    for line in [
        "（1）卷积层：通过可学习的卷积核在图像上滑动提取局部特征，具有局部连接与权值共享的特点，"
        "大幅减少参数量。给定输入尺寸 W、卷积核 K、步长 S、填充 P，输出尺寸计算见下式。",
        "（2）激活函数 ReLU：f(x)=max(0,x)，引入非线性并缓解梯度消失，是 AlexNet 的关键创新之一。",
        "（3）池化层：通过最大池化等操作进行下采样，降低特征图尺寸、增强平移不变性。",
        "（4）局部响应归一化（LRN）：对相邻通道的响应进行归一化，增强模型的泛化能力。",
        "（5）Dropout：训练时以一定概率随机置零神经元输出，有效抑制过拟合。",
        "（6）全连接层：将提取到的高层特征映射到类别空间，输出分类结果。"]:
        para(doc, line, indent=False)
    equation(doc, mr("O = ⌊") + mfrac(mr("W − K + 2P"), mr("S")) + mr("⌋ + 1"))
    para(doc, "AlexNet 由 5 个卷积层与 3 个全连接层堆叠而成，首次将 ReLU、Dropout、数据增强与 GPU 训练"
              "有机结合，是深度学习发展史上的里程碑。")

    doc.add_heading("2.4 模型评估指标", level=2)
    para(doc, "对于分类任务，依据混淆矩阵中的真正例 TP、假正例 FP、真负例 TN、假负例 FN，常用指标定义如下：")
    equation(doc, mr("准确率 = ") + mfrac(mr("TP + TN"), mr("TP + TN + FP + FN")))
    equation(doc, mr("精确率 = ") + mfrac(mr("TP"), mr("TP + FP"))
             + mr("，  召回率 = ") + mfrac(mr("TP"), mr("TP + FN")))
    equation(doc, mr("F1 = ") + mfrac(mr("2 × 精确率 × 召回率"), mr("精确率 + 召回率")))
    para(doc, "对于多分类问题，常计算各类别指标的算术平均，即宏平均（Macro-average），它对每个类别一视"
              "同仁，能更好地反映模型在不平衡数据上的综合表现。混淆矩阵则以矩阵形式展示各类别的预测分布，"
              "可直观分析模型的混淆情况。")

    doc.add_heading("2.5 实验环境配置", level=2)
    para(doc, "本实训的全部实验均在同一硬件与软件环境下完成，具体配置如表 2-1 所示。项目一的机器学习"
              "模型基于 scikit-learn 实现，项目二的 AlexNet 基于 PyTorch 实现并使用 GPU 加速训练。")
    add_table(doc, ["项目", "配置"],
              [["操作系统", "Linux (Ubuntu)"],
               ["编程语言", "Python 3.12"],
               ["深度学习框架", "PyTorch 2.7.1（CUDA 11.8）+ torchvision 0.22.1"],
               ["机器学习库", "scikit-learn 1.8.0"],
               ["数据处理", "pandas 3.0.2、numpy 2.4.3"],
               ["可视化", "matplotlib 3.10.8、seaborn"],
               ["GPU", "NVIDIA GeForce RTX 4090（24GB）"]],
              caption="表 2-1 实验环境配置")
    doc.add_page_break()


# ---------------------- 第3章 项目一 ----------------------
def chapter3(doc):
    doc.add_heading("第3章 项目一：红酒质量分类", level=1)

    doc.add_heading("3.1 数据集介绍与来源", level=2)
    para(doc, "本项目采用 UCI 机器学习仓库公开的葡萄酒质量数据集（Wine Quality Data Set），来源网址为 "
              "https://archive.ics.uci.edu/dataset/186/wine+quality 。其中白葡萄酒子集共 4898 个样本，"
              "每个样本包含 11 个理化特征以及 1 个质量评分标签（0–10 的整数，由品酒师综合评定）。"
              "各特征含义见表 3-1。")
    add_table(doc, ["特征", "含义", "特征", "含义"],
              [["fixed acidity", "固定酸度", "total sulfur dioxide", "总二氧化硫"],
               ["volatile acidity", "挥发性酸度", "density", "密度"],
               ["citric acid", "柠檬酸", "pH", "酸碱度"],
               ["residual sugar", "残糖", "sulphates", "硫酸盐"],
               ["chlorides", "氯化物", "alcohol", "酒精度"],
               ["free sulfur dioxide", "游离二氧化硫", "quality", "质量评分（标签）"]],
              caption="表 3-1 白葡萄酒数据集特征说明")

    doc.add_heading("3.2 数据加载与可视化分布", level=2)
    para(doc, "数据以分号分隔的 CSV 文件存储，使用 pandas 读入后核验为 4898×12（11 特征 + 1 标签），"
              "经检查无缺失值与重复值，数据质量良好。为深入理解数据特性，绘制了质量评分分布直方图与"
              "特征相关性热力图。")
    add_image(doc, os.path.join(T1, "eda_quality_dist.png"), 3.5, "图 3-1 葡萄酒质量评分分布")
    para(doc, "由分布图可见，质量评分高度集中于 5、6、7 三档，评分为 6 的样本最多，而极高（8、9）与极低"
              "（3、4）评分的样本极少，整体呈现明显的“中间多、两端少”的长尾不平衡特征，这给少数类的识别"
              "带来挑战。由相关性热力图可见，酒精度与质量评分呈较强正相关，而密度与质量评分呈负相关，二者"
              "本身又高度负相关，这与酿酒学常识一致，表明所选特征对质量判别具备一定有效信息。")
    add_image(doc, os.path.join(T1, "eda_corr_heatmap.png"), 5.0, "图 3-2 特征相关性热力图")

    doc.add_heading("3.3 数据预处理", level=2)
    for line in [
        "（1）标签构造：原始质量评分有 7 个取值且极不均衡，直接做七分类意义有限。故将其分箱为三个等级"
        "——评分≤5 记为“差”(类0)、=6 记为“中”(类1)、≥7 记为“好”(类2)。分箱后三类样本数分别为 "
        "1640、2198、1060，更符合实际的质量分级需求。",
        "（2）特征标准化：由于各特征量纲差异较大，采用 z=(x−μ)/σ 对 11 个特征做零均值、单位方差标准化，"
        "且仅在训练集上拟合参数后再应用到测试集，严格避免数据泄漏。",
        "（3）数据集划分：按 8:2 分层抽样划分，得训练集 3918 个、测试集 980 个样本，保证各类别比例一致。",
        "（4）不平衡处理：所有模型统一设置 class_weight='balanced'，按类别频率反比加权，使模型更关注少数类。"]:
        para(doc, line, indent=False)

    doc.add_heading("3.4 模型建立", level=2)
    para(doc, "选取支持向量机、决策树、随机森林、逻辑回归四种代表性算法（原理见 2.2 节），通过五折交叉"
              "验证结合网格搜索（评分准则为宏平均 F1）对关键超参数进行寻优。五折交叉验证将训练集均分为"
              "五份，轮流以四份训练、一份验证，取平均性能作为评分，从而更稳健地选择超参数。各模型搜索"
              "空间与最优结果见表 3-2。")
    add_table(doc, ["模型", "主要搜索空间", "最优超参数"],
              [["SVM", "C∈{1,10}, γ∈{scale,0.1}, RBF核", "C=10, γ=0.1"],
               ["决策树", "最大深度∈{None,10,20}, 叶子最小样本∈{1,5}", "深度None, 叶子1"],
               ["随机森林", "树数∈{200,400}, 最大深度∈{None,20}", "树数200, 深度20"],
               ["逻辑回归", "C∈{0.5,1,10}", "C=0.5"]],
              caption="表 3-2 四种模型的搜索空间与最优超参数")

    doc.add_heading("3.5 模型评估", level=2)
    para(doc, "在独立测试集上对四个模型进行评估，整体指标对比见表 3-3 与图 3-3。")
    rows = [["SVM", "0.617", "0.616", "0.660", "0.622"],
            ["决策树", "0.650", "0.647", "0.644", "0.645"],
            ["随机森林", "0.735", "0.748", "0.723", "0.734"],
            ["逻辑回归", "0.522", "0.524", "0.567", "0.520"]]
    csv_path = os.path.join(T1, "metrics.csv")
    if os.path.exists(csv_path):
        with open(csv_path, encoding="utf-8-sig") as f:
            data = list(csv.reader(f))
        rows = [[r[0]] + [f"{float(x):.3f}" for x in r[1:]] for r in data[1:]]
    add_table(doc, ["模型", "准确率", "精确率", "召回率", "宏平均F1"], rows,
              caption="表 3-3 四种机器学习模型测试集性能对比")
    add_image(doc, os.path.join(T1, "model_compare.png"), 5.0, "图 3-3 各模型测试集指标对比")
    para(doc, "由表 3-3 可见，随机森林在四项指标上全面领先，其次为决策树与 SVM，逻辑回归表现最差。"
              "为进一步分析最优模型，给出随机森林各类别详细指标（表 3-4）及其混淆矩阵（图 3-4）。")
    add_table(doc, ["类别", "精确率", "召回率", "F1", "样本数"],
              [["差(≤5)", "0.78", "0.73", "0.76", "328"],
               ["中(=6)", "0.69", "0.77", "0.73", "440"],
               ["好(≥7)", "0.77", "0.67", "0.72", "212"]],
              caption="表 3-4 随机森林各类别指标")
    para(doc, "从各类别指标看，随机森林对“差”“好”两端类别精确率较高，而“中”类样本最多、与相邻等级"
              "边界模糊，召回率虽高但精确率相对偏低，这与质量评分本身的连续性、相邻等级难以截然区分的"
              "特点相符。")
    add_image(doc, os.path.join(T1, "cm_随机森林.png"), 3.5, "图 3-4 随机森林混淆矩阵")
    add_image(doc, os.path.join(T1, "rf_feature_importance.png"), 4.5, "图 3-5 随机森林特征重要性排序")
    para(doc, "特征重要性结果显示酒精度、密度与挥发性酸度对质量判别贡献最大，与相关性分析的结论相互印证。")

    doc.add_heading("3.6 误差分析", level=2)
    para(doc, "进一步分析随机森林的混淆矩阵（图 3-4）可以发现，模型的误差呈现明显的“邻级混淆”特征："
              "真实为“差”的 328 个样本中有 82 个被误判为“中”，真实为“好”的 212 个样本中有 68 个被误判"
              "为“中”，而“差”与“好”两个相隔等级之间的相互误判极少（分别仅 5 例与 2 例）。其根源在于，"
              "质量评分本质上是连续的有序变量，相邻等级（如 5 分与 6 分）的样本在理化特征上高度接近，"
              "缺乏清晰的判别边界；而处于中间的“中”类样本数量最多、左右毗邻两端，自然成为误判的主要汇聚点。"
              "这提示我们：若将该问题建模为有序回归或引入对邻级错误更敏感的代价矩阵，有望进一步降低此类误差。")

    doc.add_heading("3.7 本章小结", level=2)
    para(doc, "实验表明，随机森林在所有指标上均显著领先，宏平均 F1 达 0.734，比次优的决策树高出约 9 个"
              "百分点；逻辑回归作为线性模型表现最差（F1 仅 0.520），说明该任务中特征与质量之间存在较强的"
              "非线性关系，而集成学习能更好地拟合复杂决策边界，并通过多树投票降低方差、提升鲁棒性。"
              "特征重要性进一步揭示了影响葡萄酒质量的关键理化指标，体现了机器学习模型的可解释价值。")
    doc.add_page_break()


# ---------------------- 第4章 项目二 ----------------------
def chapter4(doc):
    doc.add_heading("第4章 项目二：手写 AlexNet 图像分类", level=1)

    doc.add_heading("4.1 数据集介绍与来源", level=2)
    para(doc, "本项目采用 Fashion-MNIST 数据集，由德国 Zalando 公司发布，来源网址为 "
              "https://www.kaggle.com/datasets/zalando-research/fashionmnist （亦可经 torchvision 自动下载）。"
              "数据集共 70000 张 28×28 的灰度服饰图像，其中训练集 60000 张、测试集 10000 张，均匀覆盖 10 个"
              "类别：T恤、裤子、套衫、连衣裙、外套、凉鞋、衬衫、运动鞋、包、短靴。相比经典的手写数字 MNIST，"
              "Fashion-MNIST 类间相似度更高、纹理更复杂、识别难度更大，已成为评估图像分类算法的标准基准之一。")

    doc.add_heading("4.2 数据加载与可视化", level=2)
    para(doc, "通过 torchvision.datasets.FashionMNIST 接口自动下载并加载数据。由于 AlexNet 的标准输入为 "
              "224×224，加载时通过 Resize 将原始 28×28 图像放大至 224×224。部分测试样例及预测结果见图 4-1"
              "（绿色标题表示预测正确，红色表示预测错误），可见模型对大多数样本判别准确，仅在视觉相似的"
              "上装类别间偶有混淆。")
    add_image(doc, os.path.join(T2, "samples.png"), 5.0, "图 4-1 测试集预测样例")

    doc.add_heading("4.3 数据预处理", level=2)
    para(doc, "图像预处理流程包括：(1) Resize(224) 将图像缩放到网络所需尺寸；(2) ToTensor 将像素值转为张量"
              "并归一化到 [0,1]；(3) Normalize 按 Fashion-MNIST 的单通道均值 0.2860、标准差 0.3530 标准化，"
              "使输入分布更利于网络收敛。此外，从 60000 张训练图中随机划出 10% 作为验证集用于模型选择"
              "（保存验证集表现最优的模型），最终在独立的 10000 张测试图上评估。")

    doc.add_heading("4.4 模型建立", level=2)
    para(doc, "按照经典 AlexNet 论文，使用 nn.Conv2d、nn.MaxPool2d、nn.Linear 等基础算子逐层手工搭建网络，"
              "不调用 torchvision.models 中的任何预置 AlexNet。网络输入为 1×224×224 的灰度图，由 5 个卷积块"
              "与 3 个全连接层构成，各层配置与输出尺寸见表 4-1，总参数量约 58.3M。其中卷积层负责逐级提取由"
              "边缘、纹理到语义部件的层次化特征，池化层逐步降低空间分辨率，全连接层完成最终类别判别。")
    add_table(doc, ["层", "配置", "输出尺寸"],
              [["输入", "灰度图像", "1×224×224"],
               ["Conv1", "96核 11×11 步长4 填充2；ReLU+LRN", "96×55×55"],
               ["MaxPool1", "3×3 步长2", "96×27×27"],
               ["Conv2", "256核 5×5 填充2；ReLU+LRN", "256×27×27"],
               ["MaxPool2", "3×3 步长2", "256×13×13"],
               ["Conv3", "384核 3×3 填充1；ReLU", "384×13×13"],
               ["Conv4", "384核 3×3 填充1；ReLU", "384×13×13"],
               ["Conv5", "256核 3×3 填充1；ReLU", "256×13×13"],
               ["MaxPool3", "3×3 步长2", "256×6×6"],
               ["FC6", "Dropout(0.5)+Linear；ReLU", "4096"],
               ["FC7", "Dropout(0.5)+Linear；ReLU", "4096"],
               ["FC8", "Linear（输出层）", "10"]],
              caption="表 4-1 手写 AlexNet 网络结构")
    para(doc, "训练配置：损失函数采用交叉熵，优化器为带动量的随机梯度下降（SGD，动量 0.9，权重衰减 5e-4）：")
    equation(doc, mr("L = − Σ ") + msub(mr("y"), mr("i")) + mr(" log ") + msub(mr("ŷ"), mr("i"))
             + mr("，  v ← βv − η∇L，  θ ← θ + v"))
    para(doc, "其中动量项有助于加速收敛并抑制震荡。初始学习率 0.01，采用步长 10、衰减系数 0.1 的学习率"
              "调度（StepLR），即每 10 轮将学习率降为原来的 1/10；批大小 128，共训练 15 轮，按验证集准确率"
              "保存最优模型。训练在单张 NVIDIA RTX 4090 GPU 上完成，约耗时十余分钟。")

    doc.add_heading("4.5 模型评估", level=2)
    para(doc, "训练与验证过程的损失、准确率曲线见图 4-2。可见训练初期损失迅速下降、准确率快速提升，随后"
              "平稳收敛；验证集最高准确率为 0.9232，且在第 10 轮学习率衰减后性能出现一次小幅跃升，训练曲线"
              "与验证曲线基本贴合，未见明显过拟合，表明 Dropout 与权重衰减起到了有效的正则化作用。")
    add_image(doc, os.path.join(T2, "curves.png"), 6.0, "图 4-2 训练/验证 损失与准确率曲线")
    acc, f1 = 0.9155, 0.915
    mp = os.path.join(T2, "test_metrics.json")
    if os.path.exists(mp):
        m = json.load(open(mp)); acc, f1 = m["acc"], m["f1"]
    para(doc, f"在 10000 张测试图像上，模型最终准确率为 {acc*100:.2f}%，宏平均 F1 为 {f1:.3f}。各类别详细"
              "指标见表 4-2，整体混淆情况见图 4-3。")
    add_table(doc, ["类别", "精确率", "召回率", "F1"],
              [["T恤", "0.87", "0.85", "0.86"], ["裤子", "0.99", "0.98", "0.99"],
               ["套衫", "0.87", "0.87", "0.87"], ["连衣裙", "0.91", "0.94", "0.92"],
               ["外套", "0.84", "0.89", "0.86"], ["凉鞋", "0.98", "0.98", "0.98"],
               ["衬衫", "0.76", "0.74", "0.75"], ["运动鞋", "0.96", "0.97", "0.96"],
               ["包", "0.98", "0.98", "0.98"], ["短靴", "0.97", "0.96", "0.97"]],
              caption="表 4-2 AlexNet 各类别测试指标")
    add_image(doc, os.path.join(T2, "confusion_matrix.png"), 5.0, "图 4-3 AlexNet 测试集混淆矩阵")

    doc.add_heading("4.6 误差分析", level=2)
    para(doc, "由混淆矩阵（图 4-3）可见，AlexNet 的误差高度集中于四类上装——T恤、套衫、外套、衬衫之间。"
              "其中“衬衫”最难识别，其 1000 个测试样本中有 95 个被误判为 T恤、61 个误判为套衫、73 个误判为"
              "外套；而“T恤”也有 109 个被误判为衬衫。与之形成鲜明对比的是，“裤子”“包”“凉鞋”“短靴”等"
              "形态轮廓差异显著的类别几乎不与其他类别混淆，识别率接近完美。究其原因，这些上装在 28×28 的"
              "低分辨率灰度图中，领口、袖型、版型等关键区分性细节被严重弱化，整体轮廓高度相似，是 "
              "Fashion-MNIST 的固有难点；预测样例图（图 4-1）中的红色错误样本也大多集中于此类上装。由此可见，"
              "进一步提升性能的关键在于增强模型对细粒度纹理的判别能力，例如引入数据增强、更深的网络或注意力机制。")

    doc.add_heading("4.7 本章小结", level=2)
    para(doc, "手写 AlexNet 在 Fashion-MNIST 上取得了 91.55% 的测试准确率，验证了所实现网络的正确性与有效性。"
              "从混淆矩阵与各类指标看，“裤子”“包”“凉鞋”“短靴”等形态独特的类别识别率接近完美（F1≥0.97），"
              "而“衬衫”类最易混淆（F1 仅 0.75），主要与“T恤”“套衫”“外套”相互误判——这些上装在低分辨率"
              "灰度图下轮廓与纹理高度相似，是 Fashion-MNIST 公认的识别难点。")
    doc.add_page_break()


# ---------------------- 第5章 结论 ----------------------
def chapter5(doc):
    doc.add_heading("第5章 结论与实验总结", level=1)
    doc.add_heading("5.1 结论", level=2)
    para(doc, "（1）在结构化数据的红酒质量分类中，集成学习方法（随机森林）显著优于单一模型与线性模型，"
              "测试集宏平均 F1 达 0.734，是四种算法中的最佳选择，印证了集成学习在中小规模结构化数据上的"
              "普遍优势。", indent=False)
    para(doc, "（2）在图像分类任务中，逐层手写实现的 AlexNet 能够有效学习服饰图像的层次化特征，测试准确率"
              "达 91.55%，充分验证了深度卷积网络在视觉任务上的强大能力以及本次复现的正确性。", indent=False)
    doc.add_heading("5.2 实验总结", level=2)
    para(doc, "通过本次实训，我们完整走通了“数据加载—查看—预处理—建模—评估—结论”的机器学习全流程，"
              "并获得如下收获：其一，深入理解了传统机器学习与深度学习在数据形态、建模思路与适用场景上的"
              "差异；其二，掌握了标准化、分层划分、类别加权等预处理手段及其对结果的影响，体会到“数据决定"
              "模型上限”的重要性；其三，熟悉了网格搜索与交叉验证在超参数调优中的作用；其四，通过亲手逐层"
              "实现 AlexNet，透彻理解了卷积、池化、ReLU、LRN、Dropout 等核心组件的原理与协同机制。工程实践"
              "中还解决了绘图中文字体缺失数字字形、scikit-learn 新版本参数变更等实际问题，锻炼了独立排查与"
              "解决问题的能力。")
    doc.add_heading("5.3 改进方案", level=2)
    para(doc, "项目一：尝试 XGBoost、LightGBM 等更强的梯度提升模型；采用 SMOTE 等过采样方法更充分处理类别"
              "不平衡；引入特征工程（特征交互、分箱）与特征选择以提升判别力；也可将质量评分作为回归任务建模"
              "后再分级。")
    para(doc, "项目二：引入数据增强（随机裁剪、翻转、旋转）抑制过拟合并提升泛化；用批归一化（BatchNorm）"
              "替代 LRN 加速收敛、稳定训练；适当增加训练轮数并采用余弦退火等更精细的学习率策略；亦可对比 "
              "VGG、ResNet 等更先进结构，分析深度对精度的影响。")
    para(doc, "通用：建立更系统的实验管理与超参数搜索流程，对关键结果进行多次重复实验以评估稳定性，"
              "并引入更全面的可视化与误差分析手段。")
    doc.add_heading("参考文献", level=1)
    for ref in [
        "[1] Krizhevsky A, Sutskever I, Hinton G E. ImageNet Classification with Deep Convolutional "
        "Neural Networks. NeurIPS, 2012: 1097-1105.",
        "[2] Cortez P, Cerdeira A, Almeida F, et al. Modeling wine preferences by data mining from "
        "physicochemical properties. Decision Support Systems, 2009, 47(4): 547-553.",
        "[3] Xiao H, Rasul K, Vollgraf R. Fashion-MNIST: a Novel Image Dataset for Benchmarking "
        "Machine Learning Algorithms. arXiv:1708.07747, 2017.",
        "[4] Breiman L. Random Forests. Machine Learning, 2001, 45(1): 5-32.",
        "[5] Cortes C, Vapnik V. Support-Vector Networks. Machine Learning, 1995, 20(3): 273-297.",
        "[6] Pedregosa F, et al. Scikit-learn: Machine Learning in Python. JMLR, 2011, 12: 2825-2830.",
        "[7] Paszke A, et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library. "
        "NeurIPS, 2019: 8024-8035."]:
        para(doc, ref, indent=False)


def build():
    doc = Document()
    set_base_style(doc)
    cover(doc)
    abstract(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    out = os.path.join(HERE, "毕业论文_人工智能实训.docx")
    doc.save(out)
    print("[论文] 已生成:", out)


if __name__ == "__main__":
    build()
