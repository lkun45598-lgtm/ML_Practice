# -*- coding: utf-8 -*-
"""生成任务一、任务二两篇独立 Word 论文。

LaTeX 版分别位于：
- reports/task1_wine_quality/report.tex
- reports/task2_alexnet_fmnist/report.tex
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION


# ============ 学校论文格式辅助：页码分节 / 真实目录域 / 三线表 ============
def _add_page_field(paragraph):
    """在段落中插入 Word 的 PAGE 域（显示当前页码）。"""
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2)


def _set_pgnum(section, fmt, start=None):
    """设置该节页码格式：fmt='upperRoman'（I,II）或 'decimal'（1,2），可指定起始值。"""
    sectPr = section._sectPr
    for el in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(el)
    pg = OxmlElement("w:pgNumType"); pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def _footer_pagenum(section, roman=False):
    """该节页脚居中显示页码域。"""
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(p)


def _enable_update_fields(doc):
    """打开文档时自动更新域（让目录页码自动生成）。"""
    settings = doc.settings.element
    for el in settings.findall(qn("w:updateFields")):
        settings.remove(el)
    uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true")
    settings.append(uf)


def toc_field(doc):
    """真实目录域：标题黑体4号居中、字间空4字距；条目由 Word 生成，含小黑点引导与右对齐页码。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目" + "　　　　" + "录")  # 空4个全角字距
    r.font.name = "SimHei"
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "SimHei")
    r.font.size = Pt(14)   # 4 号字
    r.font.bold = True
    p2 = doc.add_paragraph()
    run = p2.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fsep = OxmlElement("w:fldChar"); fsep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "（请在 Word 中按 Ctrl+A 后 F9 更新目录；打开时一般自动更新）"
    fend = OxmlElement("w:fldChar"); fend.set(qn("w:fldCharType"), "end")
    for e in (f1, instr, fsep, t, fend):
        run._r.append(e)
    # 目录条目字体设为宋体小4号
    for name in ("TOC 1", "TOC 2", "TOC 3"):
        try:
            st = doc.styles[name]
            st.font.name = "SimSun"; st.font.size = Pt(12)
            st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "SimSun")
        except KeyError:
            pass


def _three_line_borders(table):
    """三线表：仅保留表格上、下边框（1.5pt）与表头下方一条线（0.75pt），去掉其余线与底纹。"""
    tblPr = table._tbl.tblPr
    for el in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(el)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        if edge in ("top", "bottom"):
            el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "12")
            el.set(qn("w:space"), "0"); el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "none"); el.set(qn("w:sz"), "0"); el.set(qn("w:space"), "0")
        borders.append(el)
    tblPr.append(borders)
    # 表头行下边线
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement("w:tcBorders")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000")
        tcB.append(b); tcPr.append(tcB)

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
T1_OUT = os.path.join(ROOT, "task1_ml_wine", "outputs")
T2_OUT = os.path.join(ROOT, "task2_alexnet_fmnist", "outputs")
REPORT_ASSETS = os.path.join(ROOT, "report")
TASK1_DIR = os.path.join(ROOT, "reports", "task1_wine_quality")
TASK2_DIR = os.path.join(ROOT, "reports", "task2_alexnet_fmnist")


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "SimSun")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    style.paragraph_format.line_spacing = 1.5


def para(doc, text, indent=True, align=None):
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    if align is not None:
        p.alignment = align
    return p


def add_center_title(doc, text, size=16):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.size = Pt(size)
    r.font.bold = True
    return p


def add_table(doc, header, rows, caption=None):
    doc.add_paragraph()  # 表与正文之间上空一行
    if caption:  # 表题在表上方居中，与表之间不留空行
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].font.size = Pt(10.5)
    table = doc.add_table(rows=1, cols=len(header))  # 不设主题样式，改用三线表
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _three_line_borders(table)
    doc.add_paragraph()  # 表与正文之间下空一行
    return table


def add_image(doc, path, width=5.5, caption=None):
    if not os.path.exists(path):
        return
    doc.add_paragraph()  # 图与正文之间上空一行
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:  # 图题在图下方居中
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10.5)
    doc.add_paragraph()  # 图与正文之间下空一行


def cover(doc, subtitle, title):
    for _ in range(3):
        doc.add_paragraph()
    add_center_title(doc, "人工智能综合实训 II", 28)
    add_center_title(doc, subtitle, 22)
    for _ in range(2):
        doc.add_paragraph()
    add_center_title(doc, title, 18)
    for _ in range(3):
        doc.add_paragraph()
    fields = [
        ("学　　院：", "人工智能与低空技术学院"),
        ("专　　业：", "人工智能"),
        ("小组成员：", "雷正（202434610309）、蔡铭飞（202434610301）、冼嘉谦（202434610326）"),
        ("指导教师：", "赵静"),
        ("提交日期：", "2026 年 6 月"),
    ]
    for k, v in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(k + v)
        r.font.size = Pt(14)
    doc.add_page_break()


def declaration(doc):
    add_center_title(doc, "华南农业大学本科毕业论文（设计）原创性声明", 14)
    para(doc, "本人郑重声明：所呈交的毕业论文（设计），是本人在导师的指导下，独立进行研究工作所取得的成果。"
              "除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的作品成果。"
              "对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的"
              "法律结果由本人承担。")
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph("作者签名：＿＿＿＿＿＿　　日期：　　年　　月　　日")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for _ in range(2):
        doc.add_paragraph()
    add_center_title(doc, "华南农业大学本科毕业论文（设计）使用授权声明", 14)
    para(doc, "本人完全了解学校有关保留、使用毕业论文（设计）的规定，同意学校保留并向国家有关部门或机构"
              "送交毕业论文（设计）的复印件和电子版，允许毕业论文（设计）被查阅和借阅。学校可以将本毕业"
              "论文（设计）的全部或部分内容编入有关数据库进行检索，可以采用影印、缩印或扫描等复制手段保存"
              "和汇编毕业论文（设计）。")
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph("作者签名：＿＿＿＿＿　指导教师签名：＿＿＿＿＿")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph("日期：　　年　　月　　日　　日期：　　年　　月　　日")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()


def task1_abstract(doc):
    add_center_title(doc, "摘　要")
    para(doc, "结构化数据分类是机器学习的典型问题。本文以支持向量机、决策树、随机森林和逻辑回归四种代表性模型为"
              "研究对象，以 UCI Wine Quality 数据集中的白葡萄酒子集为主实验平台，将原始质量评分归并为“差”“中”“好”"
              "三类，围绕数据加载、探索性分析、预处理、模型构建、模型评估和结果分析完成完整机器学习流程，并采用标准化"
              "流水线、五折交叉验证和网格搜索进行超参数选择；在此基础上进一步把同一套模型推广到难度递增的五个结构化"
              "数据集上对照，并对其中表现欠佳者做针对性处理，考察数据特点与模型选择如何共同决定性能。")
    para(doc, "实验结果表明，随机森林在白葡萄酒上取得较优表现，准确率 73.47%、宏平均 F1 为 0.7336；“回归预测再分级”"
              "的有序建模将宏平均 F1 提升至 0.7533、严重跨级误判由 7 降至 4。跨数据集对照进一步表明：不存在放之四海"
              "而皆准的最优模型，最优选择随数据集而变；类别数量并非主要难度来源；极端不平衡会使准确率严重失真。针对"
              "表现欠佳的贷款违约与臭氧数据，本文先以补充特征、SMOTE 过采样、更换更强模型等静态手段验证（宏 F1 停在约"
              "0.58—0.62），进而设计任务级创新——把臭氧还原为时间序列并注入时序特征、为贷款补全数十个真实征信特征并注入"
              "单调约束。值得强调的是，臭氧时序特征在单次划分下一度呈现 +0.058 的提升，但因该划分测试集仅含约 4 个少数类"
              "样本而高度不可靠；改用时序交叉验证汇集约 58 个样本后提升消失（Δ≈-0.003）；而为贷款补入数十个真实征信特征"
              "后宏 F1 仍仅由 0.585 变为 0.586，单调约束亦持平于约 0.57。这表明两者“效果不好”确属数据本质——即便补全真实"
              "信息也几乎无增益（利率已是征信信息的近似充分统计量）；同时给出评估教训：极端不平衡下须以交叉验证汇集足够"
              "少数类样本，方能避免单次划分的“假提升”。")
    para(doc, "关键词：机器学习；结构化数据分类；模型对照；类别不平衡；时序交叉验证；单调约束；跨数据集")
    doc.add_page_break()
    add_center_title(doc, "Abstract")
    para(doc, "Structured-data classification is a representative machine learning problem. This paper takes four representative "
              "models (SVM, decision tree, random forest, logistic regression) as its research object, with the white wine "
              "subset of the UCI Wine Quality dataset as the primary experimental platform; the original quality scores are "
              "merged into three classes (low, medium, high) and a complete workflow with standardization, 5-fold cross "
              "validation and grid search is implemented. The same models are then extended to five structured datasets of "
              "increasing difficulty, with targeted treatment applied to the under-performing ones.")
    para(doc, "Random forest performs best on white wine (accuracy 73.47%, macro-F1 0.7336); a regression-then-discretization "
              "ordinal strategy raises macro-F1 to 0.7533. The cross-dataset comparison shows there is no universally best model, "
              "the number of classes is not the main difficulty driver, and extreme imbalance severely distorts accuracy. For the "
              "weak datasets (loan default, ozone), beyond static remedies (feature enrichment, SMOTE, stronger models) we further "
              "designed task-level innovations: recasting ozone as a time series with temporal features, and adding dozens of real "
              "application-time credit features plus monotonic constraints for loan default. Notably, the ozone temporal features "
              "showed a +0.058 macro-F1 gain under a single time split, but that split contained only about 4 minority samples and "
              "was unreliable; under time-series cross-validation pooling about 58 minority samples the gain vanished (delta about "
              "-0.003); adding dozens of real credit features moved loan-default macro-F1 only from 0.585 to 0.586, and the monotonic "
              "model stayed near 0.57. This confirms their poor results stem from intrinsic data difficulty (even adding real "
              "information barely helps, since the interest rate is already a near-sufficient statistic of credit signals), and "
              "yields a methodological lesson: under extreme imbalance, cross-validation pooling enough minority samples is required "
              "to avoid illusory gains from a single split.")
    para(doc, "Keywords: machine learning; structured-data classification; model comparison; class imbalance; time-series cross-validation; monotonic constraints; cross-dataset")
    doc.add_page_break()


def task2_abstract(doc):
    add_center_title(doc, "摘　要")
    para(doc, "图像分类是深度学习中的典型任务。本文以卷积神经网络架构为研究对象，基于 PyTorch 基础模块逐层手写 "
              "AlexNet（不调用任何预置模型），并以 Fashion-MNIST 服饰图像数据集为主实验平台，完成数据加载、预处理、"
              "训练、验证、测试和误差分析的完整流程；在此基础上进一步把同一模型放到难度递增的四个数据集上对照，并与"
              "轻量 SimpleCNN 和手写小型 ResNet 比较，考察架构设计对识别性能的影响。针对 Fashion-MNIST 图像尺寸较小、"
              "类别间相似度较高等特点，本文将 28×28 灰度图像放大至 224×224，并在主模型中采用 BatchNorm、余弦退火"
              "学习率调度和数据增强等训练策略。")
    para(doc, "实验结果表明，主模型在测试集上取得 94.38% 的准确率和 0.9437 的宏平均 F1。混淆矩阵和分类样例显示，"
              "模型对裤子、包、鞋类等区分度较高的类别识别效果较好，而 T 恤、套衫、外套和衬衫等上装类别之间仍存在"
              "较明显混淆。进一步的消融实验表明，充分训练和 BatchNorm 对模型性能提升具有较大作用；复杂度对比显示，"
              "AlexNet 在低分辨率灰度图像任务上存在一定参数和计算冗余。")
    para(doc, "在此基础上，本文将手写 AlexNet 推广到难度递增的四个数据集（Fashion-MNIST、Cats vs Dogs、Flowers、"
              "Garbage），并与轻量 SimpleCNN 和手写小型 ResNet 进行同条件对照。结果表明：识别难度主要由数据特点"
              "（每类样本量、图像复杂度、类别平衡）决定，而非类别数量；在参数量几乎相同的受控对照下，3×3 小核堆叠"
              "优于大核；小型 ResNet 仅以约 1/20 参数即可追平甚至反超 AlexNet。由此得出核心结论：决定卷积网络性能的"
              "是数据特点与架构设计而非参数规模，当架构足够强时真正的瓶颈是数据量。")
    para(doc, "关键词：深度学习；卷积神经网络；AlexNet；Fashion-MNIST；图像分类；跨数据集泛化；Grad-CAM")
    doc.add_page_break()
    add_center_title(doc, "Abstract")
    para(doc, "Image classification is a representative task in deep learning. This paper takes convolutional neural network "
              "architecture as its research object: an AlexNet is built layer by layer using basic PyTorch modules (without "
              "calling any predefined model), with Fashion-MNIST as the primary experimental platform for data loading, "
              "preprocessing, training, validation, testing and error analysis. The same model is then evaluated on four "
              "datasets of increasing difficulty and compared with a lightweight SimpleCNN and a handwritten small ResNet "
              "to examine how architectural design affects recognition performance.")
    para(doc, "The experimental results show that the main model achieves a test accuracy of 94.38% and a macro-F1 score of "
              "0.9437. Further ablation experiments show that sufficient training and BatchNorm make important contributions "
              "to model performance. Complexity comparison also indicates that AlexNet has certain parameter and computation "
              "redundancy on low-resolution grayscale images.")
    para(doc, "Furthermore, the handwritten AlexNet is extended to four datasets of increasing difficulty (Fashion-MNIST, "
              "Cats vs Dogs, Flowers, Garbage) and compared under identical settings with a lightweight SimpleCNN and a "
              "handwritten small ResNet. Recognition difficulty is mainly determined by data characteristics (per-class "
              "sample size, image complexity, class balance) rather than the number of classes; under a controlled "
              "comparison with almost identical parameter counts, stacked 3x3 kernels outperform large kernels; and the "
              "small ResNet matches or surpasses AlexNet with only about 1/20 of the parameters. CNN performance is thus "
              "governed by data characteristics and architectural design rather than parameter scale.")
    para(doc, "Keywords: deep learning; convolutional neural network; AlexNet; Fashion-MNIST; image classification; cross-dataset generalization; Grad-CAM")
    doc.add_page_break()


def manual_toc(doc, entries):
    add_center_title(doc, "目　录")
    for item in entries:
        para(doc, item, indent=False)
    doc.add_page_break()


def task1_body(doc):
    doc.add_heading("第1章 引言", level=1)
    doc.add_heading("1.1 研究背景与意义", level=2)
    para(doc, "结构化数据分类是机器学习应用中最常见的任务之一，广泛存在于食品质量评价、金融风险判断、医学辅助诊断"
              "和业务决策等场景。与图像、文本等非结构化数据不同，结构化数据通常由仪器测量、业务记录或人工标注得到，"
              "具有特征含义明确、维度相对有限、样本规模中等、噪声来源可解释等特点。对于这类数据，支持向量机、"
              "决策树、随机森林、逻辑回归等传统机器学习方法仍然具有训练成本低、调试方便、结果较易解释等优势。")
    para(doc, "葡萄酒质量评价兼具实际应用价值和方法研究价值。在真实生产场景中，葡萄酒质量通常受到酸度、糖分、"
              "硫化物、密度、酒精度等多种理化因素影响，最终评分又带有一定主观性。若能够基于理化指标建立稳定的"
              "质量分级模型，可以辅助质量控制、产品筛选和生产工艺分析。Cortez 等（2009）整理的 Wine Quality 数据集将"
              "理化检测指标与质量评分对应起来，为研究化学属性与感官质量之间的统计关系提供了公开基准。")
    para(doc, "该任务的难点在于，葡萄酒质量评分并不是完全离散、边界清晰的类别标签。一方面，数据集中样本主要集中"
              "在中间分值，极高和极低评分样本较少；另一方面，质量评分具有天然顺序，相邻等级之间的误判和跨等级"
              "误判在实际含义上并不相同。因此，本研究不仅比较多种传统分类模型的性能，还进一步引入宏平均 F1、"
              "二次加权 Kappa（Quadratic Weighted Kappa）、有序平均绝对误差（Mean Absolute Error，MAE）和严重跨级误判数等指标，从误差结构角度讨论质量分级模型的适用性。")
    para(doc, "从实训和研究方法角度看，白葡萄酒质量分类任务能够覆盖结构化数据建模的完整流程，包括数据读取与查看、"
              "探索性分析、标签构造、标准化、交叉验证、模型选择、测试集评价和结果解释。通过该任务，可以较系统地"
              "理解传统机器学习在中小规模结构化数据上的建模思路，并分析不同模型在可解释性、泛化能力和错误类型上的差异。")

    doc.add_heading("1.2 国内外研究现状", level=2)
    para(doc, "围绕结构化分类任务，国内外研究已经形成了较成熟的方法体系（周志华，2016）。逻辑回归（Logistic "
              "Regression，LR）作为广义线性模型的代表，具有参数含义清晰、训练效率高等特点，常用于建立可解释的"
              "分类基线。支持向量机（Support Vector Machine，SVM）由 Cortes 和 Vapnik（1995）提出，通过最大间隔"
              "原则提高分类器泛化能力，并可借助核函数刻画非线性边界，在中小规模数据集上具有较强竞争力。")
    para(doc, "树模型及其集成方法是结构化数据建模中的重要方向。单棵决策树（Decision Tree，DT）能够形成较直观的"
              "特征划分规则，但对训练样本扰动较敏感。Breiman（2001）提出的随机森林（Random Forest，RF）通过 "
              "Bootstrap 采样、随机特征选择和多树投票降低方差，在表格数据任务中具有较好的稳定性。此后，"
              "XGBoost（Chen 和 Guestrin，2016）和 LightGBM（Ke 等，2017）等梯度提升树方法在工程效率和预测性能方面得到广泛应用。")
    para(doc, "葡萄酒质量预测方面，Cortez 等（2009）以红葡萄酒和白葡萄酒理化检测数据为基础，研究了数据挖掘方法对感官"
              "偏好建模的可行性。该数据集被广泛用于分类、回归、特征重要性分析和模型比较。已有研究通常关注理化指标"
              "与质量评分之间的统计关联，其中酒精度、挥发性酸度、密度等变量常被认为与质量判断关系较密切。")
    para(doc, "类别不平衡和有序标签是该任务中需要重点处理的两个问题。He 和 Garcia（2009）指出，不平衡学习中总体"
              "准确率可能掩盖少数类别识别效果，因此需要结合平衡准确率、F1、马修斯相关系数（Matthews Correlation "
              "Coefficient，MCC）等指标进行评价。另一方面，葡萄酒质量评分从低到高具有明确次序，将其完全视为无序"
              "类别会弱化相邻误判与严重跨级误判之间的差异。因此，本文在常规三分类实验之外补充有序回归（McCullagh，"
              "1980）分级对照，并引入二次加权 Kappa 和严重误判数，以更贴合质量分级任务的实际语义。")

    doc.add_heading("1.3 研究内容与分工", level=2)
    para(doc, "本文以四种传统机器学习模型为研究对象，以白葡萄酒质量三分类为主线，把数据集视为检验模型的试验台。"
              "主要工作包括两个层面：其一，在白葡萄酒数据上完成加载与查看、三分类标签构造、特征标准化与分层划分、"
              "四种模型训练与网格搜索、测试集指标评价、混淆矩阵与特征重要性分析，以及基于误差结构的有序建模补充实验；"
              "其二，将同一套建模管线推广到难度递增的五个结构化数据集上对照，并对其中表现欠佳的数据集做针对性处理"
              "（补充特征、不平衡重采样、更换更强模型），分析数据特点与模型选择如何共同决定性能。")
    add_table(doc, ["学号", "姓名", "主要分工"],
              [["202434610301", "蔡铭飞", "数据预处理、四种模型训练与网格搜索、跨数据集对照与不平衡/特征改进实验、结果评估"],
               ["202434610309", "雷正", "总体方案设计、实验结果分析、论文主体撰写与复现说明整理"],
               ["202434610326", "冼嘉谦", "数据可视化、图表绘制、论文排版、PPT 制作与汇报"]],
              "表 1-1 小组成员分工表")

    doc.add_heading("第2章 理论基础与实验设计", level=1)
    doc.add_heading("2.1 监督学习与多分类问题", level=2)
    para(doc, "监督学习是在带标签样本上学习从输入特征到输出标签的映射。白葡萄酒质量分类中，输入为 11 个理化"
              "检测指标，输出为由原始质量评分合并得到的“差”“中”“好”三类标签。模型训练的目标是在训练样本上"
              "降低分类损失，同时在未参与训练的测试样本上保持较好的泛化能力。")
    para(doc, "该任务的特殊性在于，三类标签虽然按照分类问题处理，但类别之间仍具有低、中、高的自然顺序。因此，"
              "评价模型时不能只看预测是否完全正确，还应关注错误是发生在相邻等级之间，还是出现“差”和“好”之间的"
              "严重跨级误判。另一方面，原始质量评分集中在中间等级，宏平均 F1、平衡准确率、MCC 和二次加权 Kappa "
              "等指标比单纯准确率更能反映模型在不同类别上的平均表现。")
    doc.add_heading("2.2 标准化、交叉验证与模型选择", level=2)
    para(doc, "白葡萄酒各理化指标的量纲差异较大，例如残糖、二氧化硫、密度和酒精度的数值范围并不相同。对于 SVM "
              "和逻辑回归等依赖距离或权重优化的模型，标准化能够避免尺度较大的特征在训练中占据不合理优势。本文将 "
              "StandardScaler 封装进 scikit-learn Pipeline，使标准化参数只在每个训练折内估计，避免交叉验证中的"
              "数据泄漏。")
    para(doc, "四种模型均使用五折交叉验证和网格搜索选择超参数，优化目标为宏平均 F1。这样设置的原因是，宏平均 F1 "
              "先分别计算每个类别的识别效果，再对类别取平均，能够降低类别样本数差异对模型选择的影响。最终模型在"
              "训练集上完成调参和拟合，测试集只用于最后一次独立评价。")
    doc.add_heading("2.3 模型方法", level=2)
    para(doc, "本文选取支持向量机、决策树、随机森林和逻辑回归四种模型进行比较。支持向量机通过最大化分类间隔提高"
              "泛化能力，并借助径向基函数（Radial Basis Function，RBF）核函数表达非线性边界；决策树通过递归划分特征空间形成分类规则，可解释性较强，"
              "但单棵树容易受到训练样本扰动影响。")
    para(doc, "随机森林在决策树基础上引入 Bootstrap 采样和随机特征子集，多棵树投票能够降低单棵树方差，因此在"
              "中小规模表格数据上通常具有较好的稳定性。逻辑回归作为线性分类基线，用于检验理化指标与质量等级之间"
              "是否主要存在简单线性关系。如果逻辑回归表现明显偏低，则说明任务更依赖非线性特征组合。四种模型的"
              "核心原理如图 2-1 所示。")
    add_image(doc, os.path.join(T1_OUT, "fig_ml_models.png"), 6.0, "图 2-1 四种传统机器学习模型原理示意：间隔最大化、规则分裂、集成投票与概率输出")
    doc.add_heading("2.4 评价指标", level=2)
    para(doc, "基础指标包括准确率、宏平均精确率、宏平均召回率和宏平均 F1。准确率反映总体预测正确比例，宏平均指标"
              "则对每个类别等权计算，更适合观察类别不均衡场景下的平均识别质量。")
    para(doc, "进阶指标包括平衡准确率、二次加权 Kappa、Cohen's Kappa、MCC 和宏平均 ROC-AUC（受试者工作特征曲线下"
              "面积，Area Under the ROC Curve）。平衡准确率本质上是"
              "各类别召回率的平均值；MCC 能综合反映混淆矩阵结构；Kappa 衡量预测与真实标签的一致性，并扣除随机一致"
              "影响；二次加权 Kappa 对跨等级误判惩罚更高，更适合葡萄酒质量这种有序标签；宏平均 ROC-AUC 则从概率"
              "排序角度补充评价模型区分能力。")

    doc.add_heading("第3章 数据集与预处理", level=1)
    doc.add_heading("3.1 数据集来源与数据查看", level=2)
    para(doc, "本文采用 UCI Wine Quality 数据集中的白葡萄酒子集，共 4898 个样本，每个样本包含 11 个理化指标和"
              "一个质量评分标签。主要特征包括固定酸度、挥发性酸度、柠檬酸、残糖、氯化物、游离二氧化硫、总二氧化硫、"
              "密度、pH 值、硫酸盐和酒精度。")
    para(doc, "这些变量均来自理化检测结果，具有明确的物理或化学含义，适合作为质量分级模型的输入特征。但理化指标"
              "不能完全替代真实品鉴过程中的香气、口感和风格偏好，因此本文的建模目标是基于可观测理化指标对质量等级"
              "进行统计预测。")
    add_image(doc, os.path.join(T1_OUT, "eda_quality_dist.png"), 4.0, "图 3-1 白葡萄酒质量评分分布")
    add_image(doc, os.path.join(T1_OUT, "eda_corr_heatmap.png"), 5.4, "图 3-2 理化指标相关性热力图")
    para(doc, "质量评分主要集中在 5、6、7 三档，极高或极低评分样本较少，说明直接预测原始整数评分会遇到少数类别"
              "样本不足的问题。相关性热力图显示，部分特征之间存在一定相关关系，例如密度与残糖、酒精度等变量存在"
              "关联，提示模型需要处理特征之间的联合影响。")
    doc.add_heading("3.2 标签构造与数据划分", level=2)
    para(doc, "实验将评分按 q≤5、q=6、q≥7 归并为“差”“中”“好”三类。该处理兼顾了类别可学习性和质量分级语义："
              "q≤5 对应低于主流评分的样本，q=6 对应中等质量样本，q≥7 对应整体表现较好的样本。")
    para(doc, "数据集按 80/20 进行分层划分，保证训练集和测试集中的类别比例基本一致。三分类标签虽然降低了原始评分"
              "的细粒度，但能够减少少数类别过稀带来的评估波动，也更符合实际质量筛选中“较差、一般、较好”的分级需求。")

    doc.add_heading("第4章 模型建立与实验结果", level=1)
    doc.add_heading("4.1 训练设置", level=2)
    para(doc, "四种模型均采用五折交叉验证和网格搜索进行超参数选择。SVM 重点搜索惩罚系数 C 和 RBF 核参数 gamma；"
              "决策树和随机森林主要搜索树深度、叶子节点样本数和树数量；逻辑回归搜索正则化强度。参数范围保持相对"
              "克制，以便重点比较模型类型和误差结构，而不是单纯追求最高分。")
    add_image(doc, os.path.join(REPORT_ASSETS, "fig_ml_pipeline.png"), 5.8, "图 4-1 任务一传统机器学习建模流程")
    add_table(doc, ["模型", "搜索空间", "最优参数"],
              [["SVM", "C∈{1,10}，gamma∈{scale,0.1}", "C=10，gamma=0.1"],
               ["决策树", "最大深度∈{None,10,20}，叶子样本∈{1,5}", "深度 None，叶子样本 1"],
               ["随机森林", "树数∈{200,400}，最大深度∈{None,20}", "树数 200，深度 20"],
               ["逻辑回归", "C∈{0.5,1,10}", "C=0.5"]],
              "表 4-1 模型超参数搜索空间与最优结果")
    add_table(doc, ["模型", "准确率", "精确率", "召回率", "宏平均 F1"],
              [["SVM", "0.617", "0.616", "0.660", "0.622"],
               ["决策树", "0.650", "0.647", "0.644", "0.645"],
               ["随机森林", "0.735", "0.748", "0.723", "0.734"],
               ["逻辑回归", "0.522", "0.524", "0.567", "0.520"]],
              "表 4-2 四种模型基础评估指标")
    para(doc, "随机森林在准确率、精确率、召回率和宏平均 F1 上均取得最高结果。逻辑回归表现最低，说明理化指标与三类"
              "质量标签之间并非简单线性可分；SVM 能表达非线性边界，但在当前参数范围和数据条件下仍低于树模型；单棵"
              "决策树相比线性基线有所提升，但泛化稳定性有限。随机森林通过多树集成降低方差，更适合处理理化指标之间"
              "的非线性组合关系。")
    add_table(doc, ["模型", "平衡准确率", "二次加权 Kappa", "MCC", "宏 ROC-AUC"],
              [["SVM", "0.660", "0.626", "0.439", "0.819"],
               ["决策树", "0.644", "0.566", "0.452", "0.729"],
               ["随机森林", "0.723", "0.720", "0.581", "0.899"],
               ["逻辑回归", "0.567", "0.491", "0.311", "0.740"]],
              "表 4-3 四种模型进阶评估指标")
    para(doc, "随机森林在平衡准确率、二次加权 Kappa、MCC 和宏平均 ROC-AUC 上也取得最高结果，说明其优势不只体现在"
              "总体准确率上。二次加权 Kappa 为 0.720，表示预测结果与真实等级具有较好有序一致性；宏 ROC-AUC 为 0.899，"
              "说明模型概率输出具有较好的类别排序能力。")
    add_image(doc, os.path.join(T1_OUT, "model_compare.png"), 5.3, "图 4-2 四种模型指标对比")
    add_image(doc, os.path.join(T1_OUT, "cm_random_forest.png"), 4.2, "图 4-3 随机森林混淆矩阵")
    add_image(doc, os.path.join(T1_OUT, "rf_feature_importance.png"), 4.8, "图 4-4 随机森林特征重要性")
    para(doc, "随机森林对“差”和“好”两端类别的精确率较高，但“中”类样本与相邻等级边界较模糊。特征重要性显示，"
              "酒精度、密度和挥发性酸度对质量等级判别贡献较大。这一结果与葡萄酒质量判断常识基本一致，但特征重要性"
              "只能作为模型解释参考，不能直接理解为严格因果关系。")

    doc.add_heading("第5章 误差分析与补充实验", level=1)
    doc.add_heading("5.1 邻级混淆现象", level=2)
    para(doc, "随机森林的主要误差集中在“差”和“中”、“中”和“好”之间，而“差”和“好”之间的严重跨级误判较少。"
              "这说明模型较难完全区分相邻等级，但整体上能够保持质量等级的排序方向。对于葡萄酒评分而言，相邻分值"
              "往往对应若干理化指标和感官因素的综合变化，边界样本在特征空间中可能存在重叠。")
    doc.add_heading("5.2 有序回归分级对照", level=2)
    para(doc, "由于葡萄酒质量评分具有天然顺序，本文补充采用“回归预测连续质量分，再按阈值分级”的方式进行对照。"
              "该方法先预测连续质量分，再映射到三类标签，因此能够保留评分从低到高的顺序结构。")
    add_table(doc, ["方法", "准确率", "宏平均 F1", "二次加权 Kappa", "严重误判数", "有序 MAE"],
              [["直接分类（随机森林）", "0.7347", "0.7336", "0.7200", "7", "0.2724"],
               ["有序回归→分级", "0.7571", "0.7533", "0.7411", "4", "0.2469"]],
              "表 5-1 直接分类与有序回归分级对比")
    add_image(doc, os.path.join(T1_OUT, "ordinal_improvement.png"), 4.0, "图 5-1 有序回归分级减少严重误判")
    para(doc, "有序回归分级使准确率由 0.7347 提升至 0.7571，宏平均 F1 由 0.7336 提升至 0.7533，严重跨级误判数由 "
              "7 降至 4。这说明在质量等级具有自然顺序的场景中，将标签有序性纳入建模和评价过程具有实际意义。该对照"
              "仍属于基础尝试，后续需要多随机种子和专门有序分类模型进一步验证稳定性。")

    doc.add_heading("第6章 跨数据集模型对照研究", level=1)
    para(doc, "前述各章均围绕白葡萄酒一个数据集展开。为进一步回答“究竟是什么决定了传统机器学习模型的分类性能”，"
              "本章以四种模型本身为研究对象，把同一套建模管线推广到难度递增的五个结构化数据集上对照，考察模型表现如何"
              "随数据特点变化。本章核心结论是：不存在放之四海而皆准的“最优模型”，性能由数据特点与“模型—数据匹配”共同决定。")
    doc.add_heading("6.1 实验设置：难度阶梯与统一模型管线", level=2)
    para(doc, "本章在白葡萄酒之外，另选取电离层雷达回波（ionosphere）、贷款违约预测（Lending Club）、干豆形态分类"
              "（Dry Bean）和臭氧层超标检测（Ozone）四个公开数据集。五个数据集在样本量、特征维度、类别数和类别平衡"
              "程度上差异明显，构成由易到难的难度阶梯，如表 6-1 所示。其中 Lending Club 原始数据约 1.19GB，本文仅取已"
              "结清贷款、剔除回款额等会泄露标签的事后特征，并按类别分层抽样 3 万条以保证可计算性。")
    add_table(doc, ["数据集", "样本数", "特征维度", "类别数", "不平衡比", "特点"],
              [["电离层 ionosphere", "351", "34", "2", "1.8", "小样本、近线性可分"],
               ["白葡萄酒 wine", "4898", "11", "3", "2.1", "有序、轻度不平衡"],
               ["贷款违约 Lending Club", "30000", "12", "2", "4.0", "大规模、弱信号"],
               ["干豆 Dry Bean", "13611", "16", "7", "6.8", "多分类、形态可分"],
               ["臭氧 Ozone", "2536", "72", "2", "33.7", "高维、极度不平衡"]],
              "表 6-1 五个结构化数据集的基本特征")
    add_image(doc, os.path.join(T1_OUT, "fig_tab_datasets.png"), 6.0, "图 6-1 五个结构化数据集的难度阶梯：从小样本近线性可分到大规模弱信号、多分类与极端不平衡")
    para(doc, "五个数据集统一采用相同的建模管线：标准化、五折交叉验证和网格搜索，并对每个数据集使用完全相同的四种模型"
              "（SVM、决策树、随机森林、逻辑回归）及超参数搜索空间，主指标为宏平均 F1，辅以准确率与平衡准确率，所有"
              "模型均开启类别权重平衡。")

    doc.add_heading("6.2 跨数据集对照结果", level=2)
    para(doc, "四种模型在五个数据集上的测试集宏平均 F1 如表 6-2 和图 6-2 所示。最优模型随数据集而改变：随机森林在电离层"
              "和白葡萄酒上最好，支持向量机在干豆上最好，逻辑回归在贷款违约上最好，而在极不平衡的臭氧上各模型表现都很"
              "差。这直观印证了“没有免费午餐”——没有任何单一模型在所有任务上占优。")
    add_table(doc, ["数据集", "SVM", "决策树", "随机森林", "逻辑回归", "最优模型"],
              [["电离层 ionosphere", "0.9393", "0.9235", "0.9541", "0.8598", "随机森林"],
               ["白葡萄酒 wine", "0.6219", "0.6449", "0.7336", "0.5204", "随机森林"],
               ["贷款违约 Lending Club", "0.5698", "0.5379", "0.5275", "0.5843", "逻辑回归"],
               ["干豆 Dry Bean", "0.9375", "0.9104", "0.9333", "0.9300", "SVM"],
               ["臭氧 Ozone", "0.5805", "0.5653", "0.4920", "0.5704", "SVM"]],
              "表 6-2 四种模型在五个数据集上的测试集宏平均 F1")
    add_image(doc, os.path.join(T1_OUT, "cross_tab_f1.png"), 6.0, "图 6-2 同一套模型在五个数据集上的宏平均 F1 对照")
    para(doc, "在特征干净、类别较可分的数据集上（电离层、干豆），四种模型差距不大、整体都较高，说明此时数据可分性决定"
              "上限、模型选择影响有限；干豆虽有 7 个类别，各模型宏平均 F1 仍达 0.91 以上，说明类别数量并非主要难度来源，"
              "与项目二结论一致。随机森林在中低维、非线性、较平衡数据上通常是稳健默认（白葡萄酒上以 0.7336 显著领先）；"
              "逻辑回归在近线性可分或弱信号数据上反而更稳（贷款违约最优）。")

    doc.add_heading("6.3 不平衡与“本质困难”：准确率的陷阱", level=2)
    para(doc, "跨数据集对照还暴露了两类被准确率掩盖的问题，如图 6-3 所示。第一类是极端类别不平衡：臭氧阳性样本仅占 2.9%，"
              "不平衡比高达 33.7。此时“全部预测为多数类”即可得到 97.05% 的准确率；随机森林测试准确率 96.85%，看似很高，"
              "却反而低于这一多数类基线，其宏平均 F1 仅 0.492、平衡准确率仅 0.499（约等于随机猜测）。这说明在极端不平衡下"
              "准确率几乎没有意义，必须以宏平均 F1 或平衡准确率衡量。")
    add_image(doc, os.path.join(T1_OUT, "cross_tab_imbalance.png"), 6.0, "图 6-3 不平衡的代价：随机森林的准确率、宏F1与平衡准确率对照")
    add_image(doc, os.path.join(T1_OUT, "fig_imbalance_trap.png"), 6.0, "图 6-4 类别不平衡下的准确率陷阱：33:1 时全预测多数类即得 97% 准确率，但宏 F1 仅约 0.5")
    para(doc, "第二类是数据本身的“本质困难”：贷款违约预测中四种模型的宏平均 F1 全部落在 0.53～0.58 的狭窄区间，无论选用"
              "何种模型都无法显著突破。这是因为可在申请时获得的特征对“是否违约”的判别信号本就较弱，性能瓶颈在数据而非"
              "模型——与项目二“当模型足够强时瓶颈转向数据”的结论相互呼应。")

    doc.add_heading("6.4 针对欠佳数据集的改进尝试与归因", level=2)
    para(doc, "对表现欠佳的贷款违约和臭氧，本文进一步尝试有针对性的处理，以判断其“效果不好”究竟源于处理不充分还是"
              "数据本质，结果如表 6-3 所示。")
    add_table(doc, ["数据集", "处理手段", "最优宏F1", "结果与归因"],
              [["贷款违约", "基线（12 维数值）", "0.5843", "—"],
               ["贷款违约", "补充申请时类别特征（37 维）", "0.5902", "几乎无提升：int_rate 已编码 grade，信息冗余"],
               ["贷款违约", "梯度提升 GBDT（37 维）", "0.5759", "更强模型亦无突破"],
               ["臭氧", "基线（class_weight 平衡）", "0.5805", "—"],
               ["臭氧", "SMOTE 过采样", "0.5644", "不升反降：高维且少数类过少，插值成噪声"],
               ["臭氧", "梯度提升 GBDT", "0.5836", "仅与基线持平，未突破"]],
              "表 6-3 两个欠佳数据集的改进尝试（测试集宏平均 F1）")
    add_image(doc, os.path.join(T1_OUT, "improve_lendingclub.png"), 5.2, "图 6-5 贷款违约：补充类别特征前后的宏平均 F1（几乎持平）")
    add_image(doc, os.path.join(T1_OUT, "improve_ozone.png"), 5.4, "图 6-6 臭氧：SMOTE 过采样与类别权重对照（SMOTE 未带来改进）")
    para(doc, "三种手段给出一致结论。其一，为贷款违约补充 grade、期限、工龄等类别特征后宏平均 F1 几乎不变，因为基线已含"
              "的利率 int_rate 实际上已编码贷款等级（经核验 A 至 G 级平均利率由 7.3% 单调升至 30.8%，高度共线），而最强的"
              "FICO 信用分在本数据集版本中并不提供。其二，对极端不平衡的臭氧采用合成少数类过采样技术（Synthetic "
              "Minority Over-sampling Technique，SMOTE）后宏平均 F1 反而下降，因为特征高达 72 维而少数类训练样本仅"
              "数十个，SMOTE 在高维空间靠极少邻居插值，合成的多为噪声。其三，换用更强的梯度提升决策树（Gradient "
              "Boosting Decision Tree，GBDT）后，两个数据集宏平均 F1 仍停在约 0.58，未能突破四个基础模型的上限。")
    para(doc, "由此判断：在当前特征、划分与参数范围内，补特征、重采样、换更强模型三类手段均未给这两个数据集带来稳定提升，"
              "表明其“效果不好”更可能源于弱判别信号或极端不平衡叠加高维小样本等数据本质，而非单一处理手段缺失（需多随机"
              "种子进一步确认）。相对更有效的是“模型与数据的匹配”：在臭氧上"
              "避开会崩溃的随机森林、改用决策树或支持向量机并配合类别权重已是较优选择。这一“诊断—多手段尝试—归因”过程，"
              "与项目二中“激进数据增强反而有害”的结论相互呼应，都说明改进必须先判断瓶颈来源、再对症下药。")

    doc.add_heading("6.5 任务级方法创新尝试与严谨评估的教训", level=2)
    para(doc, "为进一步逼问“效果不好”究竟是建模方式没贴合任务结构、还是数据信息已到上限，本节针对两个欠佳数据集各设计"
              "一项任务级方法创新。其中臭氧实验还意外给出一条关于评估严谨性的重要教训。")
    para(doc, "（一）臭氧：把“无序表格”还原为“时间序列”。经核查，臭氧数据集首列即为采集日期，覆盖 1998—2004 年逐日观测，"
              "而基线直接丢弃日期列、随机划分，等于抛弃了地面臭氧的日间自相关与季节性信息。为此构造 9 维时序特征（标签滞后"
              "1/2/3 日、近 3/7 日超标率、月份与年内日的正弦/余弦季节编码，全部仅取“过去”信息以避免泄漏），并改用按时间前"
              "80%/后 20% 划分。在该单次时序划分下，加时序特征后最优宏平均 F1 由 0.6005 升至 0.6586（+0.058）、SVM 平衡准确率"
              "由 0.855 跃升至 0.984，看似是一次漂亮的突破。")
    para(doc, "然而核查暴露该结论不可靠：按时间切分后测试窗口（2003—2004）恰好只落入约 4 个超标日，在 4 个正例上算 F1 噪声"
              "极大，足以凭运气制造“假突破”。改用时序滚动交叉验证（TimeSeriesSplit 6 折，每折在过去训练、在后续时间块测试），"
              "把各折测试预测汇集起来评估，使评估覆盖全部约 58 个超标日，结果如表 6-4 所示。")
    add_table(doc, ["评估方式", "原始 72 维", "加 9 维时序特征", "结论"],
              [["单次时序划分（测试正例仅约 4 个）", "0.6005", "0.6586（+0.058）", "表面提升，但正例过少、噪声极大"],
               ["时序CV滚动汇集（覆盖约 58 个正例）", "0.6177", "0.6149（-0.003）", "严谨评估下提升消失"]],
              "表 6-4 臭氧：单次时序划分 vs 时序CV汇集评估（最优宏平均 F1）")
    add_image(doc, os.path.join(T1_OUT, "ozone_temporal.png"), 5.8, "图 6-7 臭氧：单次时序划分下各模型宏 F1（正例极少，仅作过程展示）")
    para(doc, "在覆盖 58 个正例的稳健评估下，时序特征带来的“提升”彻底消失（Δ≈-0.003）。这说明此前 +0.058 是单次少正例划分"
              "造成的评估假象。由此得到两点结论：其一，方法学上，在极端类别不平衡下单次划分的评估高度不可靠，必须用交叉验证"
              "汇集足够多的少数类样本才能得到可信结论；其二，即便把臭氧正确建模为时间序列，宏平均 F1 仍停在约 0.61，未能突破"
              "约 0.58—0.62 平台，表明瓶颈确在数据本身的弱判别信号。")
    para(doc, "（二）贷款违约：补全真实征信特征，再注入领域先验。突破信息天花板最根本的手段是给模型更多真实信息，而非"
              "更换模型。经核查，原始 loan.csv 多达 145 列，基线却只用了 12 个数值特征——大量申请时即可获得的真实征信行为"
              "特征（距上次逾期月数、信用卡使用率、最老账户年龄、近两年严重逾期账户数、破产记录、从未逾期账户占比等数十项）"
              "被丢弃。在严格排除放款后/结果泄漏列的前提下，将其补全为 56 维并派生“信用历史长度”，以分层 5 折交叉验证对照，"
              "结果如表 6-5 所示。")
    add_table(doc, ["特征集", "最优模型", "CV 宏平均 F1"],
              [["基线 12 维数值特征", "逻辑回归", "0.5850 ± 0.006"],
               ["+ 44 项真实征信特征（共 56 维）", "逻辑回归", "0.5862 ± 0.007"]],
              "表 6-5 贷款违约：补全真实征信特征前后（分层 5 折 CV 宏平均 F1）")
    para(doc, "补入数十个真实征信特征后宏平均 F1 仅由 0.5850 变为 0.5862（Δ=+0.001），几乎没有变化。这给出一个根本而清晰"
              "的解释：利率 int_rate 本身就是放贷方用这些征信特征算出的风险定价，已是它们的近似“充分统计量”，把原始特征再"
              "加回去只是重复已被利率浓缩的信息，自然无增益。这排除了“只是没用够特征”的可能，把贷款违约的瓶颈牢牢钉在数据"
              "信息本身——唯一能突破的，是连放贷方放款时都不掌握的信息，那是任何方法都无法凭空创造的。")
    para(doc, "在确认信息侧已无余量后，再从模型先验角度注入信用评分领域知识。信用评分受监管，业界要求模型满足单调性：利率、"
              "负债收入比、历史逾期/查询/坏账越高，违约概率不得下降；年收入越高，违约概率不得上升。将这些先验作为硬约束注入"
              "GBDT（HistGradientBoosting 的 monotonic_cst），与无约束对照如表 6-6 所示。")
    add_table(doc, ["模型", "准确率", "宏平均 F1", "平衡准确率"],
              [["无约束 GBDT", "0.6395", "0.5712", "0.6243"],
               ["单调约束 GBDT（领域先验）", "0.6268", "0.5692", "0.6356"]],
              "表 6-6 贷款违约：单调约束 GBDT 与无约束对照（测试集，base 12 维）")
    para(doc, "注入领域先验后宏平均 F1 基本持平（0.5712→0.5692，Δ=-0.002），同样未突破约 0.57，印证贷款违约的“困难”确属"
              "数据本质：可用特征（利率本身已是放贷方风控定价、缺乏 FICO 强特征）判别信号有限，正确的领域约束也无法凭空"
              "创造信息——单调约束的价值在于泛化稳定性与监管可解释性，而非精度。")
    para(doc, "小结：两项任务级创新——臭氧时序建模、贷款单调约束——在严谨评估下均未突破原有平台，强一致地表明这两个数据集"
              "的瓶颈在数据信息量本身，而非模型或建模方式的缺失。这比一次“看似成功的提升”更有价值：既用对照证据支撑“数据"
              "决定上限”，也留下一条工程教训——极端不平衡下须先用交叉验证坐实结论，警惕单次划分的评估假象。")

    doc.add_heading("6.6 本章小结", level=2)
    para(doc, "本章通过五个数据集、四种模型的同条件对照得到五点结论。第一，不存在放之四海而皆准的最优模型，最优选择随"
              "数据集而变。第二，数据可分性与特征质量决定性能上限，类别数量并非主要难度来源。第三，极端类别不平衡会使"
              "准确率严重失真，必须改用宏平均 F1 或平衡准确率评价。第四，当数据信号本身较弱时性能瓶颈更可能在数据而非"
              "模型——补特征、重采样、更换更强模型乃至任务级的时序建模与单调约束创新，在严谨评估下均未带来稳定提升。第五，"
              "极端不平衡下评估方式本身至关重要：单次划分可能因少数类样本过少而产生“假提升”，须以交叉验证汇集足够样本"
              "加以坐实。综合而言，传统机器学习的性能由“数据特点”与“模型—数据匹配”共同决定，改进前须先诊断瓶颈来源。")

    doc.add_heading("第7章 实验复现说明", level=1)
    para(doc, "在项目根目录执行 python task1_ml_wine/wine_quality.py 即可复现白葡萄酒主实验；执行 python "
              "task1_ml_wine/cross_dataset_tabular.py 复现五数据集对照，python task1_ml_wine/cross_dataset_improve.py "
              "复现改进尝试；python task1_ml_wine/ozone_temporal.py 与 python task1_ml_wine/credit_monotonic.py 复现臭氧时序"
              "特征与贷款单调约束两项任务级创新。脚本会完成数据加载、标签构造、模型训练、网格搜索、评估与图表输出，结果"
              "保存至 task1_ml_wine/outputs/。五个数据集中白葡萄酒会自动下载，其余四个需预先放入 task1_ml_wine/data/。")

    doc.add_heading("第8章 结论与展望", level=1)
    para(doc, "本文完成了白葡萄酒质量三分类任务的传统机器学习建模与误差分析。随机森林在四种模型中表现较优，"
              "测试集准确率为 0.7347，宏平均 F1 为 0.7336，并在进阶指标上保持较好表现。模型主要错误集中在相邻"
              "质量等级之间，说明质量等级边界具有一定模糊性。")
    para(doc, "从方法层面看，标准化、分层划分和交叉验证是保证评价可靠性的基础；面对类别分布不均衡和标签有序性，"
              "仅报告准确率会遗漏重要误差信息；随机森林等集成模型在中小规模表格数据上具有较好的稳定性和解释辅助"
              "能力。本文已扩展到五数据集对照并比较了 GBDT 与 SMOTE；后续可进一步进行多随机种子重复(报告均值±标准差)、"
              "引入 ADASYN/欠采样/代价敏感学习与专门的有序分类模型，以更稳健地验证结论。")
    refs(doc, [
        "周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
        "Chawla N V, Bowyer K W, Hall L O, et al. SMOTE: synthetic minority over-sampling technique[J]. Journal of Artificial Intelligence Research, 2002, 16: 321-357.",
        "Chen T, Guestrin C. XGBoost: a scalable tree boosting system[C]//ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. 2016: 785-794.",
        "Cortes C, Vapnik V. Support-vector networks[J]. Machine Learning, 1995, 20(3): 273-297.",
        "Cortez P, Cerdeira A, Almeida F, et al. Modeling wine preferences by data mining from physicochemical properties[J]. Decision Support Systems, 2009, 47(4): 547-553.",
        "He H, Garcia E A. Learning from imbalanced data[J]. IEEE Transactions on Knowledge and Data Engineering, 2009, 21(9): 1263-1284.",
        "Ke G, Meng Q, Finley T, et al. LightGBM: a highly efficient gradient boosting decision tree[C]//Advances in Neural Information Processing Systems. 2017: 3146-3154.",
        "McCullagh P. Regression models for ordinal data[J]. Journal of the Royal Statistical Society: Series B, 1980, 42(2): 109-142.",
        "Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.",
    ])


def task2_body(doc):
    doc.add_heading("第1章 引言", level=1)
    doc.add_heading("1.1 研究背景与意义", level=2)
    para(doc, "图像分类是计算机视觉中的基础任务，也是深度学习方法最早取得显著突破的方向之一。卷积神经网络（Convolutional Neural Network，CNN）通过"
              "局部连接、权值共享和层次化特征提取机制，能够从图像中逐级学习边缘、纹理、局部形状和高层语义信息，"
              "相比完全连接网络更适合处理具有空间结构的视觉数据。")
    para(doc, "AlexNet 在 ImageNet 图像分类任务中的成功推动了深度卷积网络的快速发展。其结构包含多层卷积、池化、"
              "修正线性单元（Rectified Linear Unit，ReLU）激活、局部响应归一化（Local Response Normalization，LRN）、Dropout 和全连接分类器，集中体现了早期深度视觉模型的关键设计。虽然当前"
              "更深或更轻量的网络结构已经广泛应用，但 AlexNet 仍然适合作为理解卷积网络基本机制的经典模型。逐层实现"
              "AlexNet，而不是直接调用预置模型，可以更清楚地理解输入尺寸变化、卷积核设置、特征图尺寸、参数量和训练策略"
              "之间的关系。")
    para(doc, "Fashion-MNIST 是 MNIST 的替代性基准数据集，由 10 类服饰图像构成。相比手写数字图像，服饰图像类别"
              "之间存在更高的视觉相似性，特别是 T 恤、套衫、外套和衬衫等上装类别在低分辨率灰度图中边界较模糊。"
              "该数据集规模适中、类别均衡、训练成本可控，既适合验证卷积神经网络的基本分类能力，也适合分析模型在"
              "细粒度视觉差异不足时的错误类型。")
    para(doc, "本研究在 Fashion-MNIST 上逐层构建 AlexNet，并围绕训练策略、组件消融、模型复杂度和可解释性展开实验。"
              "通过批归一化（Batch Normalization，BN）、数据增强、Focal Loss、复杂度对比和梯度加权类激活映射（Gradient-weighted Class Activation Mapping，Grad-CAM）可视化，可以进一步分析模型性能来源和误差成因，"
              "避免只以单一准确率评价模型效果。")

    doc.add_heading("1.2 国内外研究现状", level=2)
    para(doc, "卷积神经网络的发展经历了从浅层结构到深层结构、从单纯提升精度到兼顾效率与可解释性的过程。LeNet（LeCun 等，"
              "1998）展示了卷积、池化和共享权重在手写字符识别中的有效性，为后续视觉网络奠定了基础。AlexNet（Krizhevsky "
              "等，2012）在大规模 ImageNet 分类任务上显著提升识别性能，证明了深层卷积网络在大规模视觉数据上的表达能力。"
              "随后，VGG（Simonyan 和 Zisserman，2015）通过堆叠小卷积核研究网络深度对性能的影响，ResNet（He 等，2016）"
              "则利用残差连接缓解深层网络训练退化问题。")
    para(doc, "在训练策略方面，早期 AlexNet 使用 ReLU 激活提升非线性表达和训练效率，并结合 Dropout（Srivastava 等，"
              "2014）缓解全连接层过拟合。批归一化（Ioffe 和 Szegedy，2015）通过标准化中间层输入分布稳定训练过程，常用于替代或弱化早期网络中的局部响应归一化。"
              "数据增强也是图像分类中常用的正则化手段，可通过随机翻转、旋转、裁剪等方式提高模型对输入扰动的鲁棒性。")
    para(doc, "在数据集方面，MNIST 长期作为图像分类入门基准，但其类别结构较简单，许多模型容易取得较高准确率。"
              "Fashion-MNIST 提供了一个与 MNIST 接口相近但更具挑战性的服饰图像基准。其上装类别形态相近，低分辨率灰度图"
              "缺乏颜色和纹理细节，因此比手写数字更适合检验模型对细粒度形状差异的表达能力。")
    para(doc, "在误差分析和模型解释方面，仅报告测试准确率难以说明卷积网络真正依赖哪些图像区域进行判断。Focal Loss（Lin 等，2017）通过"
              "降低易分样本权重，使模型更关注难分类样本；Grad-CAM（Selvaraju 等，2017）通过梯度加权的类激活映射展示模型关注区域，为分析卷积"
              "神经网络预测依据提供了可视化工具。本文在主模型训练之外引入 Focal Loss 对照和 Grad-CAM 可视化，目的在于结合"
              "混淆矩阵讨论上装类别误差来源，而不是只比较总体准确率。")

    doc.add_heading("1.3 研究内容与分工", level=2)
    para(doc, "本文以卷积神经网络架构为研究对象、以逐层手写的 AlexNet 为主模型，把数据集视为检验模型的试验台。主要工作"
              "包括两个层面：其一，以 Fashion-MNIST 十分类任务为主实验平台，完成 AlexNet 的逐层实现、训练评估和误差分析，"
              "并通过组件消融、复杂度对比、Focal Loss 对照和 Grad-CAM 可视化深入剖析模型性能来源；其二，将同一 AlexNet "
              "推广到难度递增的四个数据集，与轻量 SimpleCNN 和手写小型 ResNet 同条件对照，从卷积核大小、残差连接、分类头"
              "结构等角度分析架构设计对识别性能与泛化能力的影响。")
    add_table(doc, ["学号", "姓名", "主要分工"],
              [["202434610309", "雷正", "总体方案设计、AlexNet/ResNet/SimpleCNN 网络逐层实现、跨数据集与大核小核架构对照实验设计与训练、核心实验分析与论文主体撰写"],
               ["202434610301", "蔡铭飞", "数据加载与预处理、主模型训练与评估、组件消融与 Grad-CAM 可解释性分析、实验图表生成"],
               ["202434610326", "冼嘉谦", "数据与结果整理、图表与论文排版、PPT 制作与答辩汇报、复现说明整理"]],
              "表 1-1 小组成员分工表")

    doc.add_heading("第2章 理论基础", level=1)
    doc.add_heading("2.1 卷积神经网络", level=2)
    para(doc, "卷积神经网络通过卷积核在局部感受野上滑动提取特征。与全连接网络相比，卷积网络的核心优势在于局部"
              "连接和权值共享。局部连接符合图像中边缘、纹理和局部形状逐级组合的特点；权值共享使同一卷积核能够在"
              "整幅图像上检测相同模式，从而减少参数量并增强平移等变性。")
    para(doc, "池化层通过局部下采样降低空间分辨率并增强平移鲁棒性，ReLU 激活函数提供非线性表达能力并缓解梯度"
              "饱和问题。对于 Fashion-MNIST 这类低分辨率灰度图像，卷积网络需要在有限像素细节中提取服饰轮廓、"
              "开口形状和局部边界差异，因此卷积核尺度、池化次数和分类器容量都会影响最终识别效果。")
    doc.add_heading("2.2 AlexNet 结构", level=2)
    para(doc, "AlexNet 由五个卷积层和三个全连接层组成，原始结构包含 ReLU、局部响应归一化、Dropout 和大规模全连接"
              "分类器。本文基于 Fashion-MNIST 的单通道灰度输入，将首层输入通道由 3 改为 1，将最后一层输出类别数"
              "改为 10，并保留五个卷积层和三层全连接分类器的主体结构。")
    para(doc, "由于 AlexNet 的全连接分类器规模较大，当输入被调整为 224×224 后，最后一个卷积块输出约为 256×6×6，"
              "展平后进入 4096 维全连接层。大量参数集中在分类器部分，这有利于提高模型容量，但也可能在 Fashion-MNIST "
              "这种图像信息较简单的数据集上造成计算冗余。因此，本文后续专门进行复杂度对比。")
    add_image(doc, os.path.join(REPORT_ASSETS, "fig_alexnet_arch.png"), 5.8, "图 2-1 手写 AlexNet 网络结构示意")
    doc.add_heading("2.3 归一化、正则化与优化策略", level=2)
    para(doc, "卷积网络训练不仅依赖结构本身，也依赖归一化、正则化和优化策略。原始 AlexNet 使用 LRN 增强相邻通道"
              "之间的竞争，但后续深度学习实践中 BatchNorm 更常用于稳定中间层输入分布。本文以 BatchNorm 作为主模型"
              "配置，并保留 LRN 基线作为对照。")
    para(doc, "Dropout 是 AlexNet 全连接分类器中的重要正则化手段，可以降低隐藏单元之间的 co-adaptation，缓解大容量"
              "分类器带来的过拟合风险。优化方面，本文采用带动量的随机梯度下降（Stochastic Gradient Descent，SGD）、权重衰减和余弦退火学习率调度；数据增强则通过"
              "随机水平翻转和小角度旋转构造输入扰动，提高模型对轻微姿态变化的鲁棒性。")
    doc.add_heading("2.4 评价指标与可解释性方法", level=2)
    para(doc, "本文采用准确率、宏平均精确率、宏平均召回率和宏平均 F1 评价多分类性能。Fashion-MNIST 类别数量均衡，"
              "但不同类别难度差异较大，因此宏平均 F1 能更好反映各类别平均识别质量。ROC-AUC、Cohen's Kappa 和 MCC "
              "则从概率区分能力、一致性和相关性角度补充评价总体结果。")
    para(doc, "除数值指标外，本文使用混淆矩阵和 Grad-CAM 可视化分析错误来源。Grad-CAM 利用目标类别得分对最后一层"
              "卷积特征图的梯度生成类别激活图，可以显示模型在作出预测时更关注图像中的哪些区域。该方法用于辅助解释"
              "上装类别混淆，而不是替代定量评价。")

    doc.add_heading("第3章 数据集与模型实现", level=1)
    doc.add_heading("3.1 Fashion-MNIST 数据集", level=2)
    para(doc, "Fashion-MNIST 数据集包含 70000 张 28×28 灰度服饰图像，其中训练集 60000 张、测试集 10000 张，覆盖"
              "T 恤、裤子、套衫、连衣裙、外套、凉鞋、衬衫、运动鞋、包和短靴 10 个类别。本文从训练集中划分 10% 作为"
              "验证集，用于模型选择和训练过程监控。")
    para(doc, "该数据集类别数量相对均衡，但类别难度并不相同。裤子、包和鞋类在整体轮廓上具有明显差异，通常较容易"
              "识别；T 恤、套衫、外套和衬衫均属于上装，低分辨率灰度图中缺少材质、颜色和细节纹理信息，类别边界更"
              "模糊。因此，本文后续重点分析上装类别混淆。")
    add_image(doc, os.path.join(T2_OUT, "samples.png"), 5.4, "图 3-1 Fashion-MNIST 测试集预测样例")
    doc.add_heading("3.2 数据预处理", level=2)
    para(doc, "由于 AlexNet 标准输入尺寸为 224×224，本文通过 Resize 将 Fashion-MNIST 原图放大至 224×224。随后使用 "
              "ToTensor 转换为张量，并根据 Fashion-MNIST 的单通道均值 0.2860 和标准差 0.3530 进行归一化。")
    para(doc, "放大图像并不会产生新的细节信息，但可以使输入尺寸与 AlexNet 经典结构保持一致，避免重新设计卷积层和"
              "全连接层尺寸。其代价是计算量明显增加，因此本文在复杂度实验中讨论这种处理方式的效率问题。训练集可启用"
              "随机水平翻转和小角度旋转，验证集和测试集始终使用确定性变换，以保证评价结果稳定。")
    doc.add_heading("3.3 网络实现", level=2)
    para(doc, "本文使用 PyTorch 基础模块逐层构建 AlexNet，核心模块包括卷积层、批归一化层、ReLU 激活、最大池化、"
              "Dropout 和全连接层。实现过程中未调用预置 AlexNet 模型，而是通过基础算子手工搭建网络结构。")
    add_table(doc, ["模块", "层结构", "输出尺寸"],
              [["Conv1", "11×11 Conv, BN, ReLU, MaxPool", "96×27×27"],
               ["Conv2", "5×5 Conv, BN, ReLU, MaxPool", "256×13×13"],
               ["Conv3-5", "3×3 Conv, BN, ReLU", "256×6×6"],
               ["Classifier", "Dropout, FC4096, FC4096, FC10", "10"]],
              "表 3-1 主模型 AlexNet 结构")
    para(doc, "网络前两层使用较大卷积核和最大池化快速扩大感受野并降低空间分辨率，后三个卷积层使用 3×3 卷积进一步"
              "提取局部形状组合特征。最后一个池化层后，特征图展平进入三层全连接分类器，全连接层前后加入 Dropout "
              "以降低过拟合风险。")

    doc.add_heading("第4章 训练与评估", level=1)
    doc.add_heading("4.1 训练设置", level=2)
    para(doc, "主模型使用 SGD 优化器，动量为 0.9，权重衰减为 5e-4，初始学习率为 0.01，批大小为 128，训练 40 轮，"
              "并采用余弦退火学习率调度和数据增强。训练过程包括前向传播、交叉熵损失计算、反向传播、参数更新、验证集"
              "评估和最佳权重保存。")
    para(doc, "训练过程中使用验证集准确率监控模型状态，而不是直接根据测试集调参。40 轮训练结合余弦退火调度，使模型"
              "在前期保持较大学习率进行充分搜索，后期逐渐降低学习率以稳定收敛。")
    add_image(doc, os.path.join(REPORT_ASSETS, "fig_train_pipeline.png"), 5.8, "图 4-1 任务二 AlexNet 训练与评估流程")
    add_image(doc, os.path.join(T2_OUT, "curves.png"), 5.8, "图 4-2 主模型训练与验证曲线")
    para(doc, "从训练曲线看，主模型训练准确率和验证准确率整体随训练轮次上升，并在后期趋于稳定，说明模型已经完成较"
              "充分训练。由于训练集启用了数据增强，训练集指标与验证集指标之间存在一定差异是正常现象，不能简单理解为"
              "模型退化。")
    doc.add_heading("4.2 测试集结果", level=2)
    add_table(doc, ["准确率", "宏精确率", "宏召回率", "宏 F1", "宏 ROC-AUC", "Cohen's Kappa"],
              [["0.9438", "0.9437", "0.9438", "0.9437", "0.9972", "0.9376"]],
              "表 4-1 主模型测试集指标")
    add_image(doc, os.path.join(T2_OUT, "confusion_matrix.png"), 5.0, "图 4-3 主模型测试集混淆矩阵")
    para(doc, "主模型在测试集上取得 94.38% 的准确率和 0.9437 的宏平均 F1。宏精确率、宏召回率和宏 F1 接近，说明模型"
              "在十个类别上的平均表现较均衡。宏 ROC-AUC 达到 0.9972，表明概率输出具有较强类别区分能力；Kappa 为 "
              "0.9376，说明预测结果与真实标签之间具有较高一致性。")
    para(doc, "混淆矩阵显示，裤子、鞋类和包等类别识别效果较好，而 T 恤、套衫、外套和衬衫之间仍存在较明显混淆。"
              "这些上装类别共享相似轮廓，且图像缺少颜色、材质和细节纹理，使模型需要依赖领口、袖口、衣摆等局部线索，"
              "而这些线索在 28×28 原图中并不总是清晰。")

    doc.add_heading("第5章 对照实验与误差分析", level=1)
    doc.add_heading("5.1 组件消融实验", level=2)
    add_table(doc, ["配置", "测试准确率", "宏平均 F1", "衬衫类 F1", "较上一行"],
              [["最小配置（15轮，LRN，无增强）", "0.9155", "0.9153", "0.77", "-"],
               ["+ 充分训练（40轮，余弦退火）", "0.9314", "0.9311", "0.79", "+1.59"],
               ["+ BatchNorm 替代 LRN", "0.9432", "0.9429", "0.83", "+1.18"],
               ["+ 数据增强（主模型）", "0.9438", "0.9437", "0.83", "+0.06"]],
              "表 5-1 AlexNet 组件消融实验结果")
    add_image(doc, os.path.join(T2_OUT, "improve_compare.png"), 5.0, "图 5-1 组件消融实验结果对比")
    para(doc, "消融实验显示，15 轮 LRN 基线准确率为 0.9155，延长至 40 轮并使用余弦退火后提升至 0.9314，说明原始训练"
              "预算不足以充分发挥 AlexNet 的模型容量。BatchNorm 替代 LRN 后准确率提升至 0.9432，衬衫类 F1 也提升到"
              "约 0.83，表明归一化组件对较难类别具有明显帮助。")
    para(doc, "加入数据增强后主模型准确率为 0.9438，提升幅度较小，说明在当前增强强度和训练轮次下，增强主要起到轻微"
              "正则化作用。该结果提示，消融实验需要同时控制训练轮次、归一化方式和增强策略，否则容易把训练不充分误判为"
              "结构问题。")
    doc.add_heading("5.2 模型复杂度对比", level=2)
    add_table(doc, ["模型", "参数量", "FLOPs", "CPU 延迟", "测试准确率"],
              [["AlexNet@224", "58.302M", "1063.5M", "15.437ms", "94.38%"],
               ["SimpleCNN@28", "0.391M", "7.9M", "1.496ms", "约91.8%"]],
              "表 5-2 AlexNet 与 SimpleCNN 复杂度对比")
    add_image(doc, os.path.join(T2_OUT, "arch_compare.png"), 5.0, "图 5-2 AlexNet 与小型 ResNet 架构对照")
    para(doc, "AlexNet@224 的参数量约为 58.302M，是 SimpleCNN@28 的 149.1 倍；单图 FLOPs 约为 1063.5M，是 SimpleCNN "
              "的 134.6 倍。GPU 单图延迟差距相对较小，说明并行计算可以部分抵消大模型计算开销；但 CPU 单图延迟达到 "
              "SimpleCNN 的 10.3 倍，更能体现模型部署成本。")
    para(doc, "小型 ResNet 在测试准确率上达到 94.92%，略高于 AlexNet，同时参数量明显更小。这说明现代结构可以在更少"
              "参数下取得接近或略高性能。本文保留 AlexNet 作为主模型，是为了围绕经典卷积网络完成逐层实现、训练和分析，"
              "复杂度结果则用于客观说明其在该任务上的效率局限。")
    doc.add_heading("5.3 Focal Loss 与 Grad-CAM 分析", level=2)
    add_table(doc, ["损失函数", "测试准确率", "宏平均 F1", "衬衫类 F1", "外套类 F1"],
              [["交叉熵", "0.9247", "0.9245", "0.77", "0.88"],
               ["Focal Loss", "0.9167", "0.9163", "0.75", "0.86"]],
              "表 5-3 Focal Loss 与交叉熵对照实验")
    add_image(doc, os.path.join(T2_OUT, "gradcam.png"), 5.6, "图 5-3 Grad-CAM 可解释性可视化")
    para(doc, "Focal Loss 在当前配置下未优于交叉熵，说明上装类别错误更可能来自视觉差异不足和局部细节缺失，而不是"
              "类别样本数量不均衡。面对错误类别时，应先判断错误来源是样本不均衡、训练不足、结构容量不足，还是输入"
              "信息本身有限，再选择相应改进方法。")
    para(doc, "Grad-CAM 结果显示，模型注意力主要落在服饰主体轮廓，对区分上装所需的细粒度局部区域关注不足。这与混淆"
              "矩阵结果一致：模型能够较好区分裤子、包和鞋类等轮廓差异明显的类别，但对领口、袖长、开襟、衣摆等细节"
              "依赖较强的上装类别仍存在误判。")

    doc.add_heading("第6章 跨数据集泛化与架构对照研究", level=1)
    para(doc, "前述各章均在 Fashion-MNIST 上展开。为进一步回答“究竟是什么决定了卷积神经网络的图像分类性能”，本章把"
              "手写 AlexNet 推广到难度递增的四个图像数据集，并在相同训练条件下与轻量基线 SimpleCNN 和手写小型 ResNet "
              "进行对照。本章核心结论是：决定性能的是数据本身的特点和网络的架构设计，而非参数规模。")

    doc.add_heading("6.1 实验设置：难度阶梯与三模型对照", level=2)
    para(doc, "本章在 Fashion-MNIST 之外，另选取三个真实彩色图像数据集：Flowers（花卉，5 类）、Garbage（垃圾分类，"
              "6 类）和 Cats vs Dogs（猫狗二分类）。四个数据集在类别数、样本规模、类别平衡程度和图像复杂度上差异明显，"
              "构成由易到难的难度阶梯，其基本特征如表 6-1 所示。三个彩色数据集没有官方划分，本文统一采用按类别分层的"
              " 70%/15%/15% 划分，并对损坏图像做有效性校验后跳过。")
    add_table(doc, ["数据集", "类别数", "有效样本数", "每类均量", "不平衡比", "图像性质"],
              [["Fashion-MNIST", "10", "70000", "7000", "1.0×", "灰度、居中、单物体"],
               ["Cats vs Dogs", "2", "24998", "12499", "1.0×", "彩色照片、背景多样"],
               ["Flowers", "5", "4317", "863", "1.4×", "彩色照片、类内差异大"],
               ["Garbage", "6", "2527", "421", "4.3×", "彩色照片、类间相似且不平衡"]],
              "表 6-1 四个数据集的基本特征")
    para(doc, "对照的三个模型分别代表不同设计取向：AlexNet@224（约 58.3M 参数，大卷积核 + 大型全连接分类器）、小型 "
              "ResNet@64（约 2.8M 参数，3×3 小核 + 残差连接 + 全局平均池化，Global Average Pooling，GAP）、SimpleCNN@64（约 0.39M 参数，三层卷积的"
              "浅层基线）。为保证对照公平，三个模型均从零训练、不使用任何预训练权重，并尽量统一训练预算：均采用数据增强"
              "与余弦退火；其中 SimpleCNN 与 ResNet 各使用 3 个随机种子重复训练并报告均值与标准差。本文初期曾以较短预算"
              "运行 SimpleCNN，后将其对齐为与主模型一致的 40 轮 + 余弦退火重新训练，以避免把“训练不足”误判为“结构不行”。")

    doc.add_heading("6.2 数据特点对识别难度的决定作用", level=2)
    para(doc, "三个模型在四个数据集上的测试准确率如表 6-2 和图 6-2 所示。即便是同一个 AlexNet，其准确率也在 80% 至 95% "
              "之间大幅波动：数据充足的 Cats vs Dogs 和 Fashion-MNIST 分别达到 95.38% 和 94.38%，而样本较少的 Flowers "
              "和 Garbage 仅为 80% 左右。这说明识别难度主要由数据本身的特点决定，可归纳为三个因素。四个数据集由易到难的"
              "难度阶梯及其代表样例如图 6-1 所示。")
    add_image(doc, os.path.join(T2_OUT, "fig_datasets.png"), 5.8, "图 6-1 四个数据集的难度阶梯：从标准灰度图到复杂彩色图像")
    add_table(doc, ["数据集（类别数）", "随机基线", "AlexNet 58M", "小型 ResNet 2.8M", "SimpleCNN 0.39M"],
              [["Fashion-MNIST (10)", "0.100", "0.9438", "0.9492", "0.9181±0.002"],
               ["Cats vs Dogs (2)", "0.500", "0.9538", "0.9522±0.003", "0.9044±0.004"],
               ["Flowers (5)", "0.200", "0.8075", "0.8126±0.006", "0.7609±0.011"],
               ["Garbage (6)", "0.167", "0.8011", "0.8267±0.006", "0.7356±0.013"]],
              "表 6-2 三模型在四个数据集上的测试准确率")
    add_image(doc, os.path.join(T2_OUT, "cross_dataset_acc.png"), 5.8, "图 6-2 任务难度阶梯与三模型准确率对比")
    para(doc, "第一是每类样本量。准确率最高的两个数据集每类样本均在 7000 张以上（猫狗约 12500 张），而 Flowers、Garbage "
              "每类仅数百张。样本不足直接表现为过拟合：在同一 AlexNet 上，数据充足的猫狗训练—验证准确率差约 4.4 个百分点，"
              "而 Flowers、Garbage 分别高达 11.2 和 12.3 个百分点。同一模型的过拟合程度恰好在数据少的数据集上急剧放大，"
              "说明这是“数据量问题”而非“模型问题”。")
    para(doc, "第二是图像本身的复杂度。Fashion-MNIST 为灰度、居中、单物体且无背景，类内一致性高，因此即便有 10 个类别仍较"
              "易学习；而真实彩色照片背景杂乱、姿态光照多变、类内差异大，识别难度明显更高。")
    para(doc, "第三是类别平衡与类间相似度。Garbage 不平衡比达 4.3 倍，最小类别 trash 仅 137 张，其每类 F1 如表 6-3 所示："
              "纸张、纸板较高，而 trash 仅 0.60，叠加 glass、metal、plastic 之间本就相似，使 Garbage 的宏平均 F1（0.78）"
              "明显低于其总体准确率（0.80）——这一现象正是类别不平衡与类间混淆的典型信号。")
    add_table(doc, ["类别", "paper", "cardboard", "metal", "glass", "plastic", "trash"],
              [["F1", "0.90", "0.84", "0.80", "0.79", "0.73", "0.60"],
               ["样本数", "594", "403", "410", "501", "482", "137"]],
              "表 6-3 AlexNet 在 Garbage 数据集上的每类 F1")
    para(doc, "值得强调的是，类别数量并不是决定难度的主要因素：Fashion-MNIST 有 10 个类别却达到 94%，而 Flowers 只有 5 类"
              "反而只有 81%；猫狗虽只有 2 类，但其高准确率主要来自每类上万张的充足数据。这说明“每类有多少数据、图像是否"
              "干净、类别是否平衡”比“类别多少”更能决定识别难度。")

    doc.add_heading("6.3 大卷积核与小卷积核的受控对照", level=2)
    para(doc, "小型 ResNet（全用 3×3 小核）整体优于使用大核的 AlexNet，但二者在深度、分类头和残差连接等多方面都不同，"
              "无法据此单独归因于卷积核大小。为干净地考察核大小的影响，本文设计受控对照：在同一 AlexNet 框架内，仅把第一层"
              " 11×11 和第二层 5×5 的大核替换为感受野相近的 3×3 卷积堆叠，而通道数、各阶段输出尺寸（最终 256×6×6）、全连接"
              "分类头和训练预算全部保持一致，使两者唯一差异仅在卷积核。此时大核版与小核版总参数分别为 58.30M 与 58.55M，"
              "几乎相同。两种配置在 Flowers 和 Garbage 上各以 3 个随机种子重复训练，结果如表 6-4 所示。")
    add_table(doc, ["数据集", "卷积核配置", "测试准确率", "宏平均 F1", "过拟合 gap", "总参数"],
              [["Flowers", "大核（11×11、5×5）", "0.8126±0.018", "0.8123", "11.9%", "58.30M"],
               ["Flowers", "小核（3×3 堆叠）", "0.8297±0.008", "0.8302", "11.7%", "58.55M"],
               ["Garbage", "大核（11×11、5×5）", "0.8028±0.004", "0.7762", "14.9%", "58.30M"],
               ["Garbage", "小核（3×3 堆叠）", "0.8461±0.012", "0.8321", "13.6%", "58.55M"]],
              "表 6-4 大核与小核的受控对照（同一 AlexNet 框架，仅卷积核不同）")
    para(doc, "在参数量几乎相同的前提下，小核版在两个数据集上都取得更高准确率：Flowers 提升约 1.7 个百分点，Garbage 提升约"
              " 4.3 个百分点；宏平均 F1 的提升更大（Garbage 约 5.6 个百分点），说明小核对困难和少数类别帮助更明显，且过拟合"
              " gap 还略有下降。这在本任务上印证了 VGG 的设计思想：两个 3×3 卷积叠加即可获得与一个 5×5 相近的感受野，但参数"
              "更省、中间多一层非线性、网络也更深。性能提升并非来自更多参数，而是来自“更小的核、更深的层和更强的非线性”，"
              "且越困难的数据集收益越大。")
    add_image(doc, os.path.join(T2_OUT, "fig_kernel_compare.png"), 6.0, "图 6-3 大卷积核与小卷积核堆叠的对比：相同感受野下小核更省参、非线性更强")

    doc.add_heading("6.4 为什么小型 ResNet 更优：架构胜过参数规模", level=2)
    para(doc, "小型 ResNet 仅以约 1/20 的参数（2.8M 对 58.3M）就在四个数据集上全面追平甚至反超 AlexNet。要理解这一点，"
              "关键在于观察参数究竟分布在哪里，如表 6-5 所示。")
    add_table(doc, ["模型", "总参数", "卷积主体（提特征）", "分类头（全连接）"],
              [["AlexNet", "58.30M", "3.75M（6.4%）", "54.55M（93.6%）"],
               ["小型 ResNet", "2.78M", "2.77M（99.9%）", "0.003M（0.1%）"],
               ["SimpleCNN", "0.39M", "0.09M（23.8%）", "0.30M（76.2%）"]],
              "表 6-5 三个模型的参数分布：卷积主体与分类头")
    para(doc, "AlexNet 的 58.3M 参数中高达 93.6% 集中在三层全连接分类头，真正用于提取特征的卷积部分只有 3.75M；而小型 "
              "ResNet 用全局平均池化取代庞大全连接头，几乎 100% 的参数都用在卷积提特征上。这正是 ResNet 以极少参数取胜的"
              "根本原因——AlexNet 多出的参数大多用在了易过拟合的全连接头上，没有用在刀刃上。小型 ResNet 的优势来自三处设计："
              "其一，全部使用 3×3 小核堆叠，与上一节受控对照结论一致；其二，残差连接为梯度提供恒等捷径，使网络可堆到十余层"
              "卷积而不退化，从而学到更丰富的层级特征；其三，用全局平均池化替代大型全连接层，参数极少、自带平移不变性，"
              "过拟合风险也显著降低。")
    add_image(doc, os.path.join(T2_OUT, "fig_resblock.png"), 4.6, "图 6-4 残差块（BasicBlock）结构：捷径使输出为 F(x)+x，缓解深层退化")
    add_image(doc, os.path.join(T2_OUT, "fig_resnet_arch.png"), 6.0, "图 6-5 小型 ResNet 整体结构：Stem + 三个 Stage + 全局平均池化")
    add_image(doc, os.path.join(T2_OUT, "fig_head_compare.png"), 6.0, "图 6-6 巨型全连接头与全局平均池化头的对比：后者参数少四个数量级")
    add_image(doc, os.path.join(T2_OUT, "cross_dataset_gap.png"), 5.4, "图 6-7 AlexNet 与小型 ResNet 的过拟合程度对比")
    para(doc, "过拟合对比进一步支持上述判断：在每个数据集上，小型 ResNet 的训练—验证 gap 都不高于 AlexNet（如 Garbage 为 "
              "12.2% 对 12.3%、Cats vs Dogs 为 3.3% 对 4.4%）。同时两个模型的 gap 都在样本较少的彩色数据集上偏大（均超过 "
              "10%）。这说明当模型结构已足够强时，性能瓶颈会从“模型容量”转移到“数据量”：继续增大参数收益有限，扩充数据或"
              "加强正则化才是关键。")

    doc.add_heading("6.5 针对小数据集的优化尝试与归因分析", level=2)
    para(doc, "6.2 节指出 Flowers、Garbage 因样本少而严重过拟合（gap 超过 11%），Garbage 还存在类别不平衡（trash 仅 137 张、"
              "F1 仅 0.60）。据此本文尝试两类“对症”手段：用强数据增强（随机裁剪缩放 + 颜色抖动 + 随机擦除）抑制过拟合，用"
              "类加权交叉熵（按训练集类频倒数加权）缓解不平衡。然而将二者与标签平滑一并叠加后，Garbage 测试准确率不升反降"
              "（由 80.28% 降至 77.72%）。这一反直觉结果促使本文做因子分解以定位原因，其中训练准确率是关键诊断信号："
              "过拟合表现为训练准确率远高于验证，欠拟合则表现为训练准确率本身就很低。")
    add_table(doc, ["配置", "训练准确率", "验证准确率", "测试准确率", "宏平均 F1", "trash 类 F1"],
              [["基线（基础增强）", "95.8", "80.9", "80.3", "77.6", "0.595"],
               ["+ 仅类加权", "95.2", "81.5", "82.0", "80.6", "0.713"],
               ["+ 仅强增强", "80.0", "78.6", "78.4", "75.9", "0.581"],
               ["+ 强增强 + 类加权", "79.7", "77.4", "77.7", "75.4", "0.549"]],
              "表 6-6 Garbage 小数据优化的因子分解（AlexNet，3 种子均值）")
    para(doc, "因子分解给出清晰结论。其一，类加权是有效改进：单独施加时 Garbage 测试准确率由 80.3% 升至 82.0%，宏平均 F1 由 "
              "0.776 升至 0.806，最小类别 trash 的 F1 更从 0.595 大幅回升到 0.713，而训练准确率几乎不变（95.8% 对 95.2%），"
              "说明它没有损害特征学习，只是重新平衡了对各类别的关注。其二，强增强是性能下降的真正原因：单独施加时训练准确率"
              "由 95.8% 暴跌至 80.0%，模型陷入欠拟合，测试准确率随之降到 78.4%——随机裁剪缩放、颜色抖动等激进增强破坏了 "
              "Garbage 中区分玻璃、塑料、金属所依赖的透明度、纹理与色彩等细粒度线索。其三，二者叠加时强增强的损害盖过类加权"
              "的收益，最终表现为整体下降，这正是最初“优化失败”的来源。")
    para(doc, "对 Flowers 的观察与之呼应：强增强把训练—验证 gap 从 11.9% 压到 2.2%，过拟合被有效抑制，但测试准确率基本持平"
              "（81.3% 对 81.6%），说明其瓶颈同样不在过拟合而在样本量本身。由此得到方法学启示：面对小数据集不能盲目套用标准"
              "技巧，而应先诊断瓶颈类型再对症下药——类别不平衡可由类加权有效缓解，而过拟合在数据量受限时并非可通过激进增强"
              "消除的主要矛盾，盲目增强反而可能把过拟合“治”成欠拟合。")

    doc.add_heading("6.6 本章小结", level=2)
    para(doc, "本章通过四个数据集、三个模型的同条件对照，得到四点相互印证的结论。第一，识别难度主要由数据特点（每类样本量、"
              "图像复杂度、类别平衡）决定，而类别数量并非主要因素。第二，在参数量几乎相同的受控对照下，3×3 小核堆叠优于大核，"
              "且越困难的数据集收益越大。第三，小型 ResNet 以约 1/20 参数追平甚至反超 AlexNet，原因在于残差连接、全局平均"
              "池化和小核深网的架构设计，而非参数规模。第四，针对小数据集的优化需对症下药：类加权能有效缓解类别不平衡，而"
              "激进数据增强在小样本细粒度任务上反而会引发欠拟合。综合而言，在足够的训练配方下，决定卷积网络性能的是数据特点"
              "与架构设计；当架构足够强时，真正的瓶颈是数据量。")

    doc.add_heading("第7章 实验复现说明", level=1)
    para(doc, "在项目根目录执行 README 中的主模型训练命令可复现实验：python task2_alexnet_fmnist/experiments.py --model "
              "alexnet --bn --augment --cosine --epochs 40 --seed 0 --tag main --save-ckpt task2_alexnet_fmnist/outputs/"
              "alexnet_best.pt --save-history task2_alexnet_fmnist/outputs/history.json。随后执行 python task2_alexnet_fmnist/"
              "evaluate.py 生成测试集指标和可视化结果。")
    para(doc, "跨数据集对照实验可执行 run_cross_parallel.sh（三模型四数据集）、run_scnn_fair.sh（公平基线）、"
              "run_kernel_ablation.sh（大核 vs 小核）和 run_smalldata_opt.sh/run_smalldata_diag.sh（小数据优化及其因子分解），"
              "再用 cross_dataset_summary.py 聚合。相关开关包括 --strong-aug、--class-weight、--small-kernel 等。其中 "
              "Fashion-MNIST 会自动下载，三个彩色数据集需预先下载并解压至 task2_alexnet_fmnist/data/ 下对应目录。")

    doc.add_heading("第8章 结论与展望", level=1)
    para(doc, "本文基于 PyTorch 基础模块逐层实现 AlexNet，并在 Fashion-MNIST 十分类任务上完成训练和评估。主模型在"
              "测试集上取得 94.38% 的准确率和 0.9437 的宏平均 F1，说明 AlexNet 能够有效完成服饰图像分类任务。误差"
              "分析显示，模型主要混淆集中在视觉相似的上装类别之间。")
    para(doc, "消融实验表明，充分训练和 BatchNorm 对性能提升具有较大作用；复杂度分析说明 AlexNet 在低分辨率灰度图像"
              "任务上存在一定冗余；Grad-CAM 可视化进一步揭示模型主要关注服饰整体轮廓，对细粒度局部差异的利用仍有限。")
    para(doc, "进一步地，本文将手写 AlexNet 推广到四个难度递增的数据集，与 SimpleCNN 和小型 ResNet 同条件对照，得到三点"
              "结论：识别难度主要由数据特点决定而非类别数量；参数量几乎相同时 3×3 小核堆叠优于大核；小型 ResNet 以约 1/20 "
              "参数追平甚至反超 AlexNet。综合可见，决定卷积网络性能的是数据特点与架构设计，当架构足够强时真正的瓶颈是数据量。"
              "后续可增加多随机种子重复训练，并系统比较 VGG、ResNet、轻量 CNN 和注意力模块等结构。")
    refs(doc, [
        "He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2016.",
        "Ioffe S, Szegedy C. Batch normalization: accelerating deep network training by reducing internal covariate shift[C]//International Conference on Machine Learning. 2015.",
        "Krizhevsky A, Sutskever I, Hinton G E. ImageNet classification with deep convolutional neural networks[C]//Advances in Neural Information Processing Systems. 2012.",
        "LeCun Y, Bottou L, Bengio Y, et al. Gradient-based learning applied to document recognition[J]. Proceedings of the IEEE, 1998.",
        "Lin T Y, Goyal P, Girshick R, et al. Focal loss for dense object detection[C]//IEEE International Conference on Computer Vision. 2017.",
        "Paszke A, Gross S, Massa F, et al. PyTorch: an imperative style, high-performance deep learning library[C]//Advances in Neural Information Processing Systems. 2019.",
        "Selvaraju R R, Cogswell M, Das A, et al. Grad-CAM: visual explanations from deep networks via gradient-based localization[C]//IEEE International Conference on Computer Vision. 2017.",
        "Xiao H, Rasul K, Vollgraf R. Fashion-MNIST: a novel image dataset for benchmarking machine learning algorithms[J]. arXiv preprint arXiv:1708.07747, 2017.",
    ])


def refs(doc, items):
    doc.add_heading("参考文献", level=1)
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)


def _apply_pagination(doc):
    """正文之前（封面/声明/摘要/目录）用罗马数字，正文另起一节、页码从 1 重起。"""
    front = doc.sections[0]
    body = doc.sections[-1]
    front.different_first_page_header_footer = True  # 封面页不显示页码
    _set_pgnum(front, "upperRoman", 1)
    _footer_pagenum(front)
    _set_pgnum(body, "decimal", 1)
    _footer_pagenum(body)


def build_task1():
    doc = Document()
    set_base_style(doc)
    cover(doc, "任务一实训论文", "传统机器学习模型对结构化数据分类性能的影响研究——以白葡萄酒质量分类为主线的多数据集实证")
    declaration(doc)
    task1_abstract(doc)
    toc_field(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)  # 正文另起一节，便于页码重起
    task1_body(doc)
    _apply_pagination(doc)
    _enable_update_fields(doc)
    out = os.path.join(TASK1_DIR, "任务一_白葡萄酒质量分类论文.docx")
    doc.save(out)
    print("[任务一] 已生成:", out)


def build_task2():
    doc = Document()
    set_base_style(doc)
    cover(doc, "任务二实训论文", "卷积神经网络架构对图像分类性能的影响研究——基于 AlexNet 的多数据集实证")
    declaration(doc)
    task2_abstract(doc)
    toc_field(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)  # 正文另起一节，便于页码重起
    task2_body(doc)
    _apply_pagination(doc)
    _enable_update_fields(doc)
    out = os.path.join(TASK2_DIR, "任务二_AlexNet服饰图像分类论文.docx")
    doc.save(out)
    print("[任务二] 已生成:", out)


if __name__ == "__main__":
    build_task1()
    build_task2()
